"""Unit tests for the document ingestion layer."""

from __future__ import annotations

import codecs
import io
from pathlib import Path

import pytest

from workflow_compiler.exceptions import (
    EmptyDocumentError,
    FileValidationError,
    UnsupportedFormatError,
)
from workflow_compiler.ingestion import (
    DocumentContent,
    DocumentFormat,
    DocumentParserFactory,
    DocxParser,
    HtmlParser,
    MarkdownParser,
    PdfParser,
    SectionType,
    TextParser,
)
from workflow_compiler.ingestion.encoding import detect_and_decode

# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------


def _make_docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.core_properties.title = "Order Workflow"
    doc.core_properties.author = "QA Team"
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Customers submit orders through the portal.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Step"
    table.rows[0].cells[1].text = "Owner"
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _make_pdf_bytes() -> bytes:
    reportlab = pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.setTitle("Order Workflow PDF")
    pdf.setAuthor("QA Team")
    pdf.drawString(100, 750, "Hello PDF Document")
    pdf.drawString(100, 730, "Validate the payment details.")
    pdf.showPage()
    pdf.save()
    assert reportlab  # ensure import used
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Encoding
# ---------------------------------------------------------------------------


def test_detect_and_decode_utf8() -> None:
    text, encoding = detect_and_decode(b"plain ascii")
    assert text == "plain ascii"
    assert encoding == "utf-8"


def test_detect_and_decode_utf8_sig_strips_bom() -> None:
    data = codecs.BOM_UTF8 + b"with bom"
    text, encoding = detect_and_decode(data)
    assert text == "with bom"
    assert encoding == "utf-8-sig"


def test_detect_and_decode_utf16() -> None:
    data = "héllo wörld".encode("utf-16")  # includes BOM
    text, encoding = detect_and_decode(data)
    assert text == "héllo wörld"
    assert encoding == "utf-16"


def test_detect_and_decode_empty() -> None:
    text, encoding = detect_and_decode(b"")
    assert text == ""
    assert encoding


# ---------------------------------------------------------------------------
# TextParser
# ---------------------------------------------------------------------------


def test_text_parser_paragraph_sections() -> None:
    content = TextParser().parse(b"First paragraph.\n\nSecond paragraph.")
    assert isinstance(content, DocumentContent)
    assert content.document_format is DocumentFormat.TXT
    paragraphs = content.sections_of_type(SectionType.PARAGRAPH)
    assert [s.text for s in paragraphs] == ["First paragraph.", "Second paragraph."]
    assert content.metadata.word_count == 4
    assert content.metadata.char_count == len(content.text)
    assert content.metadata.encoding == "utf-8"


def test_text_parser_accepts_str_content() -> None:
    content = TextParser().parse("just text")
    assert content.text == "just text"
    assert content.metadata.size_bytes == len(b"just text")


# ---------------------------------------------------------------------------
# MarkdownParser
# ---------------------------------------------------------------------------


def test_markdown_parser_structure_and_title() -> None:
    md = (
        "# Order Intake\n\n"
        "Some **bold** and a [link](http://x).\n\n"
        "- first\n- second\n\n"
        "```\ncode line\n```\n"
    )
    content = MarkdownParser().parse(md.encode(), filename="doc.md")
    assert content.metadata.title == "Order Intake"
    headings = content.sections_of_type(SectionType.HEADING)
    assert headings and headings[0].level == 1
    assert len(content.sections_of_type(SectionType.LIST)) == 2
    assert content.sections_of_type(SectionType.CODE)
    # Inline markdown stripped from plain text.
    assert "bold" in content.text
    assert "**" not in content.text
    assert "http://x" not in content.text


# ---------------------------------------------------------------------------
# HtmlParser
# ---------------------------------------------------------------------------


def test_html_parser_metadata_and_sections() -> None:
    html = (
        "<html><head><title>Workflow</title>"
        '<meta name="author" content="Jane Doe"></head>'
        "<body><h1>Heading</h1><p>A paragraph.</p>"
        "<script>console.log('x')</script></body></html>"
    )
    content = HtmlParser().parse(html.encode(), content_type="text/html")
    assert content.metadata.title == "Workflow"
    assert content.metadata.author == "Jane Doe"
    assert "console.log" not in content.text
    headings = content.sections_of_type(SectionType.HEADING)
    assert headings and headings[0].text == "Heading"
    assert content.sections_of_type(SectionType.PARAGRAPH)


# ---------------------------------------------------------------------------
# DocxParser
# ---------------------------------------------------------------------------


def test_docx_parser_full(tmp_path: Path) -> None:
    path = tmp_path / "workflow.docx"
    path.write_bytes(_make_docx_bytes())

    content = DocxParser().parse_file(path)
    assert content.document_format is DocumentFormat.DOCX
    assert content.metadata.title == "Order Workflow"
    assert content.metadata.author == "QA Team"
    assert "Customers submit orders" in content.text
    assert content.sections_of_type(SectionType.HEADING)
    assert content.sections_of_type(SectionType.TABLE)


def test_docx_parser_rejects_str_content() -> None:
    with pytest.raises(FileValidationError):
        DocxParser().parse("not a real path or bytes")


def test_docx_parser_rejects_corrupt_bytes() -> None:
    from workflow_compiler.exceptions import ParseError

    with pytest.raises(ParseError):
        DocxParser().parse(b"not a docx at all")


# ---------------------------------------------------------------------------
# PdfParser
# ---------------------------------------------------------------------------


def test_pdf_parser_full() -> None:
    content = PdfParser().parse(_make_pdf_bytes(), filename="workflow.pdf")
    assert content.document_format is DocumentFormat.PDF
    assert "Hello PDF Document" in content.text
    assert content.metadata.page_count == 1
    assert content.sections_of_type(SectionType.PAGE)


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------


def test_empty_document_rejected() -> None:
    with pytest.raises(EmptyDocumentError):
        TextParser().parse(b"")


def test_oversize_document_rejected() -> None:
    parser = TextParser(max_size_bytes=8)
    with pytest.raises(FileValidationError):
        parser.parse(b"way too many bytes here")


def test_missing_file_rejected(tmp_path: Path) -> None:
    with pytest.raises(FileValidationError):
        TextParser().parse_file(tmp_path / "does_not_exist.txt")


def test_descriptor_mismatch_warns() -> None:
    content = TextParser().parse(b"hello", filename="report.pdf")
    assert content.warnings


# ---------------------------------------------------------------------------
# Factory
# ---------------------------------------------------------------------------


@pytest.fixture
def factory() -> DocumentParserFactory:
    return DocumentParserFactory()


def test_factory_supported_formats(factory: DocumentParserFactory) -> None:
    assert factory.supported_formats == frozenset(DocumentFormat)


def test_factory_for_format(factory: DocumentParserFactory) -> None:
    assert isinstance(factory.for_format(DocumentFormat.DOCX), DocxParser)


def test_factory_select_by_extension(factory: DocumentParserFactory) -> None:
    parser = factory.select(filename="notes.md")
    assert isinstance(parser, MarkdownParser)


def test_factory_select_by_content_type(factory: DocumentParserFactory) -> None:
    parser = factory.select(content_type="text/html; charset=utf-8")
    assert isinstance(parser, HtmlParser)


def test_factory_parse_file_end_to_end(
    factory: DocumentParserFactory, tmp_path: Path
) -> None:
    path = tmp_path / "doc.txt"
    path.write_text("Activate fulfillment.", encoding="utf-8")
    content = factory.parse(path)
    assert content.document_format is DocumentFormat.TXT
    assert content.metadata.filename == "doc.txt"
    assert "Activate fulfillment." in content.text


def test_factory_unsupported_extension(factory: DocumentParserFactory) -> None:
    with pytest.raises(UnsupportedFormatError):
        factory.select(filename="archive.xyz")


def test_factory_requires_a_hint(factory: DocumentParserFactory) -> None:
    with pytest.raises(UnsupportedFormatError):
        factory.select()
