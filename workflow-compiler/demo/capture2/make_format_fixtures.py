"""Generate real .docx/.pdf/.html/.txt renderings of the reference workflow document,
then check the ingestion layer actually recovers the load-bearing content from each.

Why this exists: `tests/test_ingestion.py` covers the four parsers, but only against
toy inputs (a two-line PDF, a three-paragraph DOCX). That proves the parsers run; it
does not prove a *real* workflow document survives them well enough for the pipeline
to extract facts from it. `docs/PIPELINE_HANDOFF.md` §0 lists these formats as
untested for exactly that reason.

The check is content-level, not byte-level: each format is a genuine rendering (Word
paragraphs, a laid-out PDF, semantic HTML), so the assertion is that the *substance*
the compiler depends on comes back — both workflow families, the cross-workflow
identifier, every named exception, and the compensation pairs. Those are the fields
§5 of the handoff flags as most at risk of being silently lost.

Usage:
    python make_format_fixtures.py [--out-dir DIR]

Exits non-zero if any format loses required content.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

# The document is authored once, structurally, and rendered into every format, so
# the four fixtures are genuinely the same document rather than four hand-written
# files that could drift apart.
TITLE = "Customer Lifecycle Operations"
INTRO = (
    "This document describes two related but distinct business processes operated "
    "by the Customer Platform team: customer onboarding and account provisioning. "
    "Provisioning consumes the customer record produced by onboarding."
)

# (heading, [paragraph | bullet-list]) where a list is a tuple of lines.
SECTIONS: list[tuple[str, list[str | tuple[str, ...]]]] = [
    (
        "Onboarding Purpose",
        [
            "The customer onboarding workflow registers a new customer: it validates "
            "the application, verifies the customer's identity, creates the customer "
            "record, and notifies the customer of the outcome."
        ],
    ),
    (
        "Onboarding Trigger",
        [
            "The workflow starts when a customer application is submitted and an "
            "application.received request reaches the Onboarding Service."
        ],
    ),
    (
        "Onboarding Inputs and Outputs",
        [
            "Inputs",
            (
                "application_id - identifier of the submitted application",
                "email - the applicant's email address",
            ),
            "Outputs",
            (
                "customer_record_id - identifier of the created customer record",
                "onboarding_status - registered or rejected",
            ),
        ],
    ),
    (
        "Onboarding Process",
        [
            (
                "1. The Onboarding Service validates the application using "
                "application_id and returns whether the application is complete.",
                "2. If the application is incomplete, the workflow raises "
                "ApplicationIncomplete and rejects the application; if the "
                "application is complete, the workflow continues.",
                "3. The Identity Service verifies the customer identity for email "
                "and returns a verification_id.",
                "4. The Customer Service creates the customer record and returns a "
                "customer_record_id.",
                "5. The Notification Service notifies the customer of the "
                "registration outcome.",
            )
        ],
    ),
    (
        "Onboarding Error Handling",
        [
            (
                "ApplicationIncomplete: the application is missing required fields "
                "-> reject the application and end the workflow.",
                "IdentityCheckFailed: identity verification fails -> reject the "
                "application and end the workflow.",
            )
        ],
    ),
    (
        "Onboarding Retries",
        [
            (
                "Verify the customer identity: retry up to 3 times with exponential "
                "backoff starting at 2 seconds.",
                "Notify the customer: retry up to 5 times; failure is non-fatal.",
            )
        ],
    ),
    (
        "Provisioning Purpose",
        [
            "The account provisioning workflow prepares a registered customer's "
            "account for use: it reserves an account number, configures the account, "
            "and activates it. If activation fails after configuration, the "
            "configuration is rolled back."
        ],
    ),
    (
        "Provisioning Trigger",
        [
            "The workflow starts when a customer record is registered and an "
            "account.provision request reaches the Provisioning Service. It requires "
            "the customer_record_id produced by customer onboarding."
        ],
    ),
    (
        "Provisioning Inputs and Outputs",
        [
            "Inputs",
            (
                "customer_record_id - identifier of the registered customer record",
                "plan_code - the subscription plan to provision",
            ),
            "Outputs",
            (
                "account_id - identifier of the activated account",
                "provisioning_status - active or failed",
            ),
        ],
    ),
    (
        "Provisioning Process",
        [
            (
                "1. The Provisioning Service reserves an account number for "
                "customer_record_id and returns an account_id.",
                "2. The Provisioning Service configures the account for plan_code.",
                "3. If the configuration is invalid, the workflow raises "
                "ConfigurationInvalid; if the configuration is valid, the workflow "
                "continues.",
                "4. The Provisioning Service activates the account and returns the "
                "final provisioning_status.",
            )
        ],
    ),
    (
        "Provisioning Error Handling",
        [
            (
                "ConfigurationInvalid: the plan configuration cannot be applied -> "
                "roll back the provisioning (release the account number) and end the "
                "workflow.",
                "ActivationFailed: activation does not complete -> roll back the "
                "provisioning (deconfigure the account, release the account number).",
            )
        ],
    ),
    (
        "Provisioning Retries",
        [
            (
                "Configure the account: retry up to 3 times with exponential backoff "
                "starting at 1 second.",
            )
        ],
    ),
    (
        "Provisioning Compensation",
        [
            (
                "Release the account number compensates Reserves an account number - "
                "return the reserved number if provisioning is rolled back.",
                "Deconfigure the account compensates Configures the account - undo "
                "the plan configuration if provisioning is rolled back after "
                "configuration.",
            )
        ],
    ),
]

# Substance the compiler depends on. If any of these does not survive a format, the
# downstream extraction cannot be correct no matter how good the model is.
REQUIRED = [
    "Onboarding",
    "Provisioning",
    "customer_record_id",  # the cross-workflow dependency
    "application_id",
    "plan_code",
    "ApplicationIncomplete",
    "IdentityCheckFailed",
    "ConfigurationInvalid",
    "ActivationFailed",
    "Release the account number",  # compensations — §5's highest-risk field
    "Deconfigure the account",
    "3 times",
    "5 times",
]


def write_txt(path: Path) -> None:
    lines = [TITLE, "", INTRO, ""]
    for heading, blocks in SECTIONS:
        lines += [heading, ""]
        for block in blocks:
            if isinstance(block, tuple):
                lines += [f"- {item}" for item in block]
            else:
                lines.append(block)
            lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_html(path: Path) -> None:
    parts = [
        "<!doctype html>",
        '<html lang="en"><head><meta charset="utf-8">',
        f"<title>{TITLE}</title>",
        '<meta name="author" content="Customer Platform team">',
        "</head><body>",
        f"<h1>{TITLE}</h1>",
        f"<p>{INTRO}</p>",
    ]
    for heading, blocks in SECTIONS:
        parts.append(f"<h2>{heading}</h2>")
        for block in blocks:
            if isinstance(block, tuple):
                parts.append("<ul>")
                parts += [f"<li>{item}</li>" for item in block]
                parts.append("</ul>")
            else:
                parts.append(f"<p>{block}</p>")
    parts.append("</body></html>")
    path.write_text("\n".join(parts), encoding="utf-8")


def write_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.core_properties.title = TITLE
    doc.core_properties.author = "Customer Platform team"
    doc.add_heading(TITLE, level=1)
    doc.add_paragraph(INTRO)
    for heading, blocks in SECTIONS:
        doc.add_heading(heading, level=2)
        for block in blocks:
            if isinstance(block, tuple):
                for item in block:
                    doc.add_paragraph(item, style="List Bullet")
            else:
                doc.add_paragraph(block)
    doc.save(str(path))


def write_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate

    # Real flowed text with automatic pagination — the realistic case. A canvas
    # drawString fixture (what the unit tests use) never exercises line wrapping,
    # which is precisely where PDF extraction tends to mangle identifiers.
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(path), pagesize=LETTER, title=TITLE, author="Customer Platform team"
    )
    story: list[object] = [Paragraph(TITLE, styles["Title"]), Paragraph(INTRO, styles["BodyText"])]
    for heading, blocks in SECTIONS:
        story.append(Paragraph(heading, styles["Heading2"]))
        for block in blocks:
            if isinstance(block, tuple):
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(i, styles["BodyText"])) for i in block],
                        bulletType="bullet",
                    )
                )
            else:
                story.append(Paragraph(block, styles["BodyText"]))
    doc.build(story)


WRITERS = {".txt": write_txt, ".html": write_html, ".docx": write_docx, ".pdf": write_pdf}


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out-dir", default="fixtures-formats")
    args = ap.parse_args()

    out = Path(args.out_dir)
    out.mkdir(parents=True, exist_ok=True)

    from workflow_compiler.ingestion import DocumentParserFactory

    factory = DocumentParserFactory()
    failed = False

    for ext, writer in WRITERS.items():
        path = out / f"customer_lifecycle{ext}"
        writer(path)

        content = factory.parse(path)
        text = content.text or ""
        # Normalise whitespace: a wrapped PDF line breaks mid-sentence, which is a
        # rendering artefact rather than lost content.
        flat = " ".join(text.split())
        missing = [m for m in REQUIRED if m not in flat]

        status = "OK  " if not missing else "FAIL"
        if missing:
            failed = True
        print(
            f"{status} {path.name:32s} {path.stat().st_size:>7d}B  "
            f"format={content.metadata.document_format.value:8s} "
            f"chars={len(text):>6d} sections={len(content.sections):>3d} "
            f"title={content.metadata.title!r}"
        )
        if missing:
            print(f"       MISSING: {missing}")

    print("\nall formats recovered the required content" if not failed else "\nFAILURES above")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
