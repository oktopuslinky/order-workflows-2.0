"""A small styled writer over python-docx reproducing the reference document look.

Conventions (measured on the manager's ``Business_Docs/*.docx``):

* Letter page, 0.75 in margins. The reference styles carry **no** font defaults,
  so Word renders them in its fallback — Times New Roman 10 pt — and that is
  what the writer sets explicitly (``BODY_FONT`` / ``BODY_PT``) so the export
  looks the same next to the originals.
* Title paragraph: Normal style, bold, 22 pt, colour ``2F5496``; subtitle 14 pt,
  colour ``444444``; a thin ``AAAAAA`` rule; then ``Label: value`` lines (label
  bold); another rule.
* Word built-in *Heading 1* / *Heading 2* (colour ``2E74B5``, 16 / 13 pt) —
  *Heading 3* (``1F4D78``, 12 pt) is used only for the TDD's Existing / Proposed
  parts.
* Bullets: *List Paragraph* style + a ``•`` bullet numbering level (``1.`` lists
  use real decimal numbering); the EPIC statement is a left-barred callout.
* Tables: full width (9500 dxa), single ``auto`` borders, header row shaded
  ``2F5496`` with white bold text and ``tblHeader``, body cells ``FFFFFF``,
  cell margins 80/100 dxa.
* Checklists: plain paragraphs indented 360 dxa starting with ``☑  `` / ``☐  ``.
* Inline code: Consolas runs coloured ``AA3377``; code blocks: one paragraph with
  ``CCCCCC`` borders, ``F5F5F5`` shading and Consolas 9.5 pt, lines split by
  soft breaks.

Everything is deterministic — the writer never looks at the clock (``created``
core property is fixed) so identical input yields identical bytes.
"""

from __future__ import annotations

import io
import re
from collections.abc import Iterable, Sequence
from dataclasses import dataclass
from datetime import UTC, datetime

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Emu, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .package import stabilise_package

TITLE_COLOR = "2F5496"
SUBTITLE_COLOR = "444444"
HEADING_COLOR = "2E74B5"
HEADER_FILL = "2F5496"
BODY_FILL = "FFFFFF"
CODE_COLOR = "AA3377"
CODE_FONT = "Consolas"
BODY_FONT = "Times New Roman"
BODY_PT = 10
RULE_COLOR = "AAAAAA"
CHECKED = "☑  "
UNCHECKED = "☐  "
TABLE_WIDTH_DXA = 9500
_FIXED_CREATED = datetime(2026, 1, 1, tzinfo=UTC)

_INLINE = re.compile(r"(`[^`]+`|\*\*[^*]+?\*\*|(?<![*\w])\*[^*\s][^*]*?\*(?![*\w]))")


@dataclass(frozen=True)
class Span:
    """One formatted piece of inline text."""

    text: str
    code: bool = False
    bold: bool = False
    italic: bool = False


def parse_inline(text: str, *, bold: bool = False, italic: bool = False) -> list[Span]:
    """Split ``text`` into spans: `` `code` ``, ``**bold**``, ``*italic*`` and plain.

    Deterministic and forgiving: unbalanced markers stay literal text.
    """
    spans: list[Span] = []
    pos = 0
    for match in _INLINE.finditer(text):
        if match.start() > pos:
            spans.append(Span(text[pos : match.start()], bold=bold, italic=italic))
        token = match.group(0)
        if token.startswith("`"):
            spans.append(Span(token[1:-1], code=True, bold=bold, italic=italic))
        elif token.startswith("**"):
            spans.append(Span(token[2:-2], bold=True, italic=italic))
        else:
            spans.append(Span(token[1:-1], bold=bold, italic=True))
        pos = match.end()
    if pos < len(text):
        spans.append(Span(text[pos:], bold=bold, italic=italic))
    return [s for s in spans if s.text]


def plain_text(text: str) -> str:
    """``text`` with inline markers removed (for headings and titles)."""
    return "".join(s.text for s in parse_inline(text))


def _shade(element: BaseOxmlElement, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _spacing(paragraph: Paragraph, *, before: int | None = None, after: int | None = None) -> None:
    fmt = paragraph.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before / 20)
    if after is not None:
        fmt.space_after = Pt(after / 20)


class DocxWriter:
    """Build one document in the reference style; ``bytes()`` serialises it."""

    def __init__(self) -> None:
        self.doc: DocumentType = Document()
        self._bullet_num_id: int | None = None
        self._numbered_num_id: int | None = None
        self._setup()

    # ------------------------------------------------------------------ setup
    def _setup(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Emu(12240 * 635)
        section.page_height = Emu(15840 * 635)
        margin = Emu(1080 * 635)
        section.left_margin = section.right_margin = margin
        section.top_margin = section.bottom_margin = margin
        normal = self.doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(BODY_PT)
        for name, size, color in (
            ("Heading 1", 16, HEADING_COLOR),
            ("Heading 2", 13, HEADING_COLOR),
            ("Heading 3", 12, "1F4D78"),
        ):
            style = self.doc.styles[name]
            style.font.color.rgb = RGBColor.from_string(color)
            style.font.size = Pt(size)
            style.font.name = BODY_FONT
            style.font.bold = False
        # python-docx's template maps the theme fonts to Calibri; pin the East-Asian /
        # complex-script slots too so every run really falls back to the body font.
        rpr = self.doc.styles["Normal"].element.get_or_add_rPr()
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            fonts.set(qn(attr), BODY_FONT)
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
            fonts.attrib.pop(qn(attr), None)
        props = self.doc.core_properties
        props.author = "workflow-compiler"
        props.last_modified_by = "workflow-compiler"
        props.created = _FIXED_CREATED
        props.modified = _FIXED_CREATED
        props.revision = 1

    # -------------------------------------------------------------- primitives
    def title(self, text: str) -> Paragraph:
        p = self.doc.add_paragraph()
        run = p.add_run(plain_text(text))
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor.from_string(TITLE_COLOR)
        _spacing(p, after=200)
        return p

    def subtitle(self, text: str) -> Paragraph:
        p = self.doc.add_paragraph()
        run = p.add_run(plain_text(text))
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(SUBTITLE_COLOR)
        _spacing(p, after=300)
        return p

    def blank(self) -> Paragraph:
        return self.doc.add_paragraph()

    def rule(self) -> Paragraph:
        """The thin grey line the reference puts around the metadata block."""
        p = self.doc.add_paragraph()
        ppr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), RULE_COLOR)
        border.append(bottom)
        ppr.append(border)
        _spacing(p, after=240)
        return p

    def meta(self, label: str, value: str) -> Paragraph:
        p = self.doc.add_paragraph()
        lead = p.add_run(f"{label}: ")
        lead.bold = True
        self._add_spans(p, parse_inline(value))
        _spacing(p, after=60)
        return p

    def heading(self, text: str, level: int = 1) -> Paragraph:
        level = max(1, min(level, 3))
        p = self.doc.add_paragraph(style=f"Heading {level}")
        p.add_run(plain_text(text))
        _spacing(p, before=320 if level == 1 else 260, after=160 if level == 1 else 120)
        return p

    def paragraph(
        self, text: str, *, italic: bool = False, spans: Sequence[Span] | None = None
    ) -> Paragraph:
        p = self.doc.add_paragraph()
        self._add_spans(p, spans if spans is not None else parse_inline(text, italic=italic))
        _spacing(p, after=160)
        return p

    def note(self, text: str) -> Paragraph:
        """A quoted/coverage note: italic, grey."""
        p = self.doc.add_paragraph()
        for run in self._add_spans(p, parse_inline(text, italic=True)):
            run.font.color.rgb = RGBColor.from_string(SUBTITLE_COLOR)
        _spacing(p, after=160)
        return p

    def bullet(self, text: str, *, level: int = 0) -> Paragraph:
        p = self.doc.add_paragraph(style="List Paragraph")
        self._number(p, self._bullet_num(), level)
        self._add_spans(p, parse_inline(text))
        _spacing(p, after=80)
        return p

    def numbered(self, text: str, number: int) -> Paragraph:
        """A ``1.`` list item with real decimal numbering; ``number == 1`` starts a new list."""
        if number == 1 or self._numbered_num_id is None:
            self._numbered_num_id = self._new_num("decimal")
        p = self.doc.add_paragraph(style="List Paragraph")
        self._number(p, self._numbered_num_id, 0)
        self._add_spans(p, parse_inline(text))
        _spacing(p, after=80)
        return p

    def callout(self, text: str) -> Paragraph:
        """An indented statement with a thick ``2F5496`` left bar (the EPIC statement)."""
        p = self.doc.add_paragraph()
        ppr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), TITLE_COLOR)
        border.append(left)
        ppr.append(border)
        p.paragraph_format.left_indent = Emu(400 * 635)
        self._add_spans(p, parse_inline(text))
        _spacing(p, after=200)
        return p

    def checklist_item(self, text: str, done: bool = False) -> Paragraph:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Emu(360 * 635)
        p.add_run(CHECKED if done else UNCHECKED)
        self._add_spans(p, parse_inline(text))
        _spacing(p, after=80)
        return p

    def code_block(self, code: str) -> Paragraph:
        p = self.doc.add_paragraph()
        ppr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        for side in ("top", "bottom", "left", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "4")
            el.set(qn("w:color"), "CCCCCC")
            border.append(el)
        ppr.append(border)
        _shade(ppr, "F5F5F5")
        _spacing(p, before=120, after=240)
        lines = code.replace("\r\n", "\n").rstrip("\n").split("\n")
        for i, line in enumerate(lines):
            run = p.add_run(line)
            self._monospace(run, size=9.5)
            if i < len(lines) - 1:
                run.add_break(WD_BREAK.LINE)
        return p

    def table(
        self,
        columns: Sequence[str],
        rows: Iterable[Sequence[str]],
        *,
        widths: Sequence[int] | None = None,
    ) -> None:
        """A bordered table with the shaded header row; cell text may contain ``\\n``."""
        rows = [list(r) for r in rows]
        table = self.doc.add_table(rows=1 + len(rows), cols=len(columns))
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
        tbl_pr.append(tbl_w)
        borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), "auto")
            borders.append(el)
        tbl_pr.append(borders)
        if widths is None:
            widths = self._even_widths(len(columns))
        header = table.rows[0]
        tr_pr = header._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:tblHeader"))
        for i, column in enumerate(columns):
            cell = header.cells[i]
            self._cell_props(cell, widths[i], HEADER_FILL)
            p = cell.paragraphs[0]
            run = p.add_run(plain_text(column))
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
        for r, row in enumerate(rows, start=1):
            values = list(row) + [""] * (len(columns) - len(row))
            for i in range(len(columns)):
                cell = table.rows[r].cells[i]
                self._cell_props(cell, widths[i], BODY_FILL)
                self._cell_text(cell, str(values[i]))
        # The reference documents leave an empty paragraph after each table.
        self.doc.add_paragraph()

    def sources(self, entries: Iterable[tuple[str, str]]) -> None:
        """Bulleted ``path — spans`` lines with the path in Consolas."""
        for path, detail in entries:
            p = self.doc.add_paragraph(style="List Paragraph")
            self._number(p, self._bullet_num(), 0)
            self._monospace(p.add_run(path))
            if detail:
                p.add_run(f" — {detail}")
            _spacing(p, after=60)

    def bytes(self) -> bytes:
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return stabilise_package(buffer.getvalue())

    # ---------------------------------------------------------------- helpers
    def _add_spans(self, paragraph: Paragraph, spans: Iterable[Span]) -> list[Run]:
        runs: list[Run] = []
        for span in spans:
            run = paragraph.add_run(span.text)
            if span.code:
                self._monospace(run)
            if span.bold:
                run.bold = True
            if span.italic:
                run.italic = True
            runs.append(run)
        return runs

    @staticmethod
    def _monospace(run: Run, *, size: float | None = None) -> None:
        run.font.name = CODE_FONT
        rpr = run._r.get_or_add_rPr()
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            fonts.set(qn(attr), CODE_FONT)
        if size is not None:
            run.font.size = Pt(size)
        else:
            run.font.color.rgb = RGBColor.from_string(CODE_COLOR)

    @staticmethod
    def _even_widths(n: int) -> list[int]:
        base = TABLE_WIDTH_DXA // n
        widths = [base] * n
        widths[-1] += TABLE_WIDTH_DXA - base * n
        return widths

    def _cell_props(self, cell: _Cell, width: int, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:type"), "dxa")
        tc_w.set(qn("w:w"), str(width))
        tc_pr.append(tc_w)
        _shade(tc_pr, fill)
        margins = OxmlElement("w:tcMar")
        for side, w in (("top", 80), ("left", 100), ("bottom", 80), ("right", 100)):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:type"), "dxa")
            el.set(qn("w:w"), str(w))
            margins.append(el)
        tc_pr.append(margins)

    def _cell_text(self, cell: _Cell, text: str) -> None:
        lines = text.replace("\r\n", "\n").split("\n")
        first = cell.paragraphs[0]
        for i, line in enumerate(lines):
            p = first if i == 0 else cell.add_paragraph()
            self._add_spans(p, parse_inline(line))

    @staticmethod
    def _number(paragraph: Paragraph, num_id: int, level: int) -> None:
        ppr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num)
        ppr.append(num_pr)

    def _bullet_num(self) -> int:
        if self._bullet_num_id is None:
            self._bullet_num_id = self._new_num("bullet")
        return self._bullet_num_id

    def _new_num(self, kind: str) -> int:
        """Add a numbering definition (``bullet``: • / ○ / ■, ``decimal``: 1. 2. 3.)."""
        numbering = self.doc.part.numbering_part.element
        abstract_ids = [
            int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))
        ]
        num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
        abstract_id = max(abstract_ids, default=0) + 1
        num_id = max(num_ids, default=0) + 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "hybridMultilevel")
        abstract.append(multi)
        glyphs = ("•", "○", "■") if kind == "bullet" else ("%1.", "%2.", "%3.")
        for lvl_index, glyph in enumerate(glyphs):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(lvl_index))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            fmt = OxmlElement("w:numFmt")
            fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
            text = OxmlElement("w:lvlText")
            text.set(qn("w:val"), glyph)
            jc = OxmlElement("w:lvlJc")
            jc.set(qn("w:val"), "left")
            ppr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(460 + 360 * lvl_index))
            ind.set(qn("w:hanging"), "260")
            ppr.append(ind)
            for el in (start, fmt, text, jc, ppr):
                lvl.append(el)
            abstract.append(lvl)
        # abstractNum elements must precede num elements.
        nums = numbering.findall(qn("w:num"))
        if nums:
            nums[0].addprevious(abstract)
        else:
            numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
        numbering.append(num)
        return num_id
