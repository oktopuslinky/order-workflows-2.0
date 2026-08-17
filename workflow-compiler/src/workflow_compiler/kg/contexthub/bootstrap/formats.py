"""Format-agnostic text extraction for any indexable repo file.

No filename or path heuristics — detection is by extension and file structure only.
"""
from __future__ import annotations

import csv
import io
import json
from pathlib import Path

MAX_FILE_BYTES = 1_000_000
# Keep enrichment prompts comfortably below small local-model context limits.
LLM_INPUT_CHARS = 2_800

INDEXABLE_DOC_EXTENSIONS = {
    ".md", ".rst", ".adoc", ".txt", ".csv", ".tsv", ".yaml", ".yml",
    ".json", ".mmd", ".mermaid", ".docx", ".xlsx", ".xls", ".html", ".htm", ".xml",
    ".pdf",  # workflow-compiler edit: routed through workflow_compiler.ingestion
}


def is_indexable_document(path: Path) -> bool:
    return path.suffix.lower() in INDEXABLE_DOC_EXTENSIONS


def extract_text(path: Path, *, max_bytes: int = MAX_FILE_BYTES) -> str:
    """Return plain text for indexing, retrieval, and LLM enrichment."""
    try:
        if not path.is_file() or path.stat().st_size > max_bytes:
            return ""
    except OSError:
        return ""

    ext = path.suffix.lower()
    if ext in {".md", ".rst", ".adoc", ".txt", ".mmd", ".mermaid", ".html", ".htm", ".xml"}:
        return _read_utf8(path)
    if ext in {".csv", ".tsv"}:
        return _extract_csv(path, delimiter="\t" if ext == ".tsv" else ",")
    if ext in {".yaml", ".yml"}:
        return _extract_yaml(path)
    if ext == ".json":
        return _extract_json(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in {".xlsx", ".xls"}:
        return _extract_xlsx(path)
    if ext == ".pdf":
        return _extract_pdf(path)
    return _read_utf8(path)


def _extract_pdf(path: Path) -> str:
    """workflow-compiler edit: PDFs go through the host app's document parsers."""
    try:
        from workflow_compiler.ingestion import DocumentParserFactory
    except Exception:  # pragma: no cover - host app always present
        return ""
    try:
        parsed = DocumentParserFactory().parse(path.read_bytes(), filename=path.name)
    except Exception:
        return ""
    return parsed.text or ""


def clip_for_llm(text: str, limit: int = LLM_INPUT_CHARS) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "\n… (truncated for LLM)"


def _read_utf8(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        return ""


def _extract_csv(path: Path, delimiter: str = ",") -> str:
    raw = _read_utf8(path)
    if not raw.strip():
        return ""
    out: list[str] = []
    try:
        for row in csv.reader(io.StringIO(raw), delimiter=delimiter):
            if any(cell.strip() for cell in row):
                out.append(" | ".join(cell.strip() for cell in row))
    except csv.Error:
        return raw
    return "\n".join(out)


def _extract_yaml(path: Path) -> str:
    raw = _read_utf8(path)
    if not raw.strip():
        return ""
    try:
        import yaml
        data = yaml.safe_load(raw)
    except Exception:
        return raw
    if data is None:
        return raw
    try:
        return yaml.dump(data, default_flow_style=False, allow_unicode=True)
    except Exception:
        return raw


def _extract_json(path: Path) -> str:
    raw = _read_utf8(path)
    if not raw.strip():
        return ""
    try:
        data = json.loads(raw)
        return json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return raw


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        doc = Document(path)
    except Exception:
        return ""
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ""
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return ""
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"## Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)
