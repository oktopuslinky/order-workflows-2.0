"""HTTP tests for the change-request routes: create → wizard jobs → artifacts → approve."""

from __future__ import annotations

import time
from collections.abc import Iterator
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from workflow_compiler.api.app import app
from workflow_compiler.api.auth import get_user_store
from workflow_compiler.api.dependencies import get_change_service, get_kg_service
from workflow_compiler.change.service import ChangeRequestService
from workflow_compiler.kg import InMemoryKnowledgeBaseStore, KgService
from workflow_compiler.kg.ingest import zip_folder
from workflow_compiler.storage.change_store import InMemoryChangeRequestStore
from workflow_compiler.storage.user_store import InMemoryUserStore

from .test_change_wizard import BCR_TEXT, ScriptedAnalyst
from .test_kg_service import build_corpus


@pytest.fixture
def analyst() -> ScriptedAnalyst:
    return ScriptedAnalyst()


@pytest.fixture
def client(tmp_path: Path, analyst: ScriptedAnalyst) -> Iterator[tuple[TestClient, str]]:
    kg = KgService(InMemoryKnowledgeBaseStore(tmp_path / "state"), lambda n, m: analyst)
    changes = ChangeRequestService(InMemoryChangeRequestStore(), kg, lambda n, m: analyst)
    app.dependency_overrides[get_kg_service] = lambda: kg
    app.dependency_overrides[get_change_service] = lambda: changes
    users = InMemoryUserStore()
    app.dependency_overrides[get_user_store] = lambda: users
    with TestClient(app) as test_client:
        test_client.post(
            "/auth/register",
            json={"email": "cr@example.com", "password": "password123", "display_name": "CR"},
        )
        upload = test_client.post(
            "/knowledge-bases",
            files={
                "file": (
                    "kb_mini.zip",
                    zip_folder(build_corpus(tmp_path / "kb_mini")),
                    "application/zip",
                )
            },
            data={"name": "Mini KB", "enrich": "false"},
        )
        assert upload.status_code == 202, upload.text
        _poll_job(test_client, upload.json()["job"]["job_id"])
        yield test_client, upload.json()["kb_id"]
    app.dependency_overrides.clear()


def _poll_job(client: TestClient, job_id: str, timeout: float = 30.0) -> dict:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        body = client.get(f"/jobs/{job_id}").json()
        if body["status"] != "running":
            return body
        time.sleep(0.05)
    raise AssertionError("job did not finish")


def _create(client: TestClient, kb_id: str, **extra: str) -> dict:
    response = client.post(
        "/change-requests",
        files={"file": ("BCR-001-partial-shipment.md", BCR_TEXT.encode("utf-8"), "text/markdown")},
        data={"kb_id": kb_id, **extra},
    )
    assert response.status_code == 201, response.text
    return response.json()


def test_create_list_get_delete(client: tuple[TestClient, str]) -> None:
    c, kb_id = client
    body = _create(c, kb_id, provider="nemotron")
    cr = body["change_request"]
    assert cr["kb_id"] == kb_id and cr["kb_name"] == "Mini KB"
    assert cr["bcr_meta"]["doc_id"] == "BCR-001"
    assert [r["id"] for r in cr["requirements"]] == ["BCR-01-01", "BCR-01-02", "BCR-01-03"]
    assert cr["wizard"]["provider"] == "nemotron"
    assert body["current_step"] == "impact" and body["question"] is None and body["job"] is None

    listing = c.get("/change-requests").json()["change_requests"]
    assert [item["cr_id"] for item in listing] == [cr["cr_id"]]
    assert listing[0]["current_step"] == "impact" and listing[0]["stage"] == "created"

    got = c.get(f"/change-requests/{cr['cr_id']}").json()
    assert got["change_request"]["title"] == cr["title"]
    assert c.get("/change-requests/nope").status_code == 404
    assert c.get("/change-requests/../x").status_code in (404, 422)

    assert c.delete(f"/change-requests/{cr['cr_id']}").json()["status"] == "deleted"
    assert c.get("/change-requests").json()["change_requests"] == []


def test_create_validation_errors(client: tuple[TestClient, str]) -> None:
    c, kb_id = client
    assert c.post("/change-requests", data={"kb_id": kb_id}).status_code == 400  # no doc
    assert c.post("/change-requests", data={"kb_id": "missing", "text": "x"}).status_code == 404
    bad = c.post("/change-requests", data={"kb_id": kb_id, "text": "x", "provider": "gpt"})
    assert bad.status_code == 422
    unsupported = c.post(
        "/change-requests",
        files={"file": ("bcr.exe", b"MZ", "application/octet-stream")},
        data={"kb_id": kb_id},
    )
    assert unsupported.status_code == 415


def test_wizard_flow_over_http(client: tuple[TestClient, str], analyst: ScriptedAnalyst) -> None:
    c, kb_id = client
    cr_id = _create(c, kb_id)["change_request"]["cr_id"]
    base = f"/change-requests/{cr_id}"

    # start → sync init + cr_questions job
    started = c.post(f"{base}/wizard/start", json={})
    assert started.status_code == 202, started.text
    body = started.json()
    assert body["change_request"]["wizard"]["started_at"] is not None
    assert body["change_request"]["ids"]["epic_id"] == "EPIC-001"
    job = body["job"]
    assert job["kind"] == "cr_questions" and job["scope_kind"] == "change_request"
    assert job["scope_id"] == cr_id
    done = _poll_job(c, job["job_id"])
    assert done["status"] == "succeeded", done
    # a second start while asking is idempotent (no new job)
    again = c.post(f"{base}/wizard/start", json={}).json()
    assert again["job"] is None
    assert again["question"] == "Consolidated or itemized invoice?"
    assert again["question_options"][0]["label"] == "One consolidated invoice"
    wizard = c.get(f"{base}/wizard").json()
    assert wizard["change_request"]["wizard"]["steps"][0]["status"] == "asking"

    # jobs list is filterable by scope
    jobs = c.get(f"/jobs?scope_id={cr_id}&scope_kind=change_request").json()
    assert {j["kind"] for j in jobs} == {"cr_questions"}

    # answer with follow-up, then resolve; skip the second
    ans = c.post(f"{base}/wizard/answer", json={"answer": "Not sure"})
    assert ans.status_code == 200, ans.text
    assert ans.json()["question"] == "Per order or per group?"
    ans = c.post(f"{base}/wizard/answer", json={"answer": "Per order", "option": "Per order"})
    assert ans.json()["question"] == "Cancel a single group?"
    skipped = c.post(f"{base}/wizard/skip").json()
    assert skipped["question"] is None
    assert c.post(f"{base}/wizard/skip").status_code == 409

    # draft (job) → artifact v1 with sources
    drafted = c.post(f"{base}/wizard/draft", json={"step": "impact"})
    assert drafted.status_code == 202, drafted.text
    assert drafted.json()["job"]["kind"] == "cr_draft"
    assert (
        c.post(f"{base}/wizard/draft", json={"step": "impact"}).status_code == 409
    )  # one at a time
    done = _poll_job(c, drafted.json()["job"]["job_id"])
    assert done["status"] == "succeeded", done
    art = c.get(f"{base}/artifacts/impact").json()
    assert art["version"] == 1 and art["status"] == "drafted"
    assert art["markdown"].startswith("# Impact Analysis — BCR-001")
    assert art["sources"] and art["history"][0]["source"] == "llm_draft"
    assert c.get(f"{base}/artifacts/impact?version=7").status_code == 400
    assert c.get(f"{base}/artifacts/nope").status_code == 404
    ahead = c.post(f"{base}/wizard/draft", json={"step": "tdd"})  # accepted, but the engine refuses
    assert ahead.status_code == 202
    assert _poll_job(c, ahead.json()["job"]["job_id"])["status"] == "failed"

    # human edit → v2 (human_edit); version lookup; bad edit → 400
    edited = c.put(
        f"{base}/artifacts/impact",
        json={
            "markdown": art["markdown"].replace("Structural change", "Big change"),
            "note": "wording",
        },
    )
    assert edited.status_code == 200, edited.text
    assert edited.json()["version"] == 2 and edited.json()["history"][-1]["source"] == "human_edit"
    v1 = c.get(f"{base}/artifacts/impact?version=1").json()
    assert v1["requested_version"] == 1 and "Structural change" in v1["markdown"]
    assert c.put(f"{base}/artifacts/impact", json={"markdown": "no title"}).status_code == 400

    # revise (job) → v3 llm_revision
    revised = c.post(f"{base}/wizard/revise", json={"step": "impact", "message": "Add a section"})
    assert revised.status_code == 202, revised.text
    assert _poll_job(c, revised.json()["job"]["job_id"])["status"] == "succeeded"
    art = c.get(f"{base}/artifacts/impact").json()
    assert art["version"] == 3 and art["history"][-1]["source"] == "llm_revision"
    assert c.post(f"{base}/wizard/revise", json={"step": "epic", "message": "x"}).status_code == 409

    # approve → cursor advances and the next step's questions job starts
    approved = c.post(f"{base}/artifacts/impact/approve")
    assert approved.status_code == 200, approved.text
    body = approved.json()
    assert body["current_step"] == "epic"
    assert body["change_request"]["artifacts"]["impact"]["status"] == "approved"
    assert body["job"]["kind"] == "cr_questions"
    assert _poll_job(c, body["job"]["job_id"])["status"] == "succeeded"
    listing = c.get("/change-requests").json()["change_requests"][0]
    assert listing["current_step"] == "epic" and listing["stage"] == "in_progress"

    # cannot approve an undrafted step
    assert c.post(f"{base}/artifacts/epic/approve").status_code == 409


def test_other_users_projects_are_visible_when_shared(client: tuple[TestClient, str]) -> None:
    c, kb_id = client
    cr_id = _create(c, kb_id)["change_request"]["cr_id"]
    c.post("/auth/logout")
    c.post(
        "/auth/register",
        json={"email": "other@example.com", "password": "password123", "display_name": "O"},
    )
    assert c.get(f"/change-requests/{cr_id}").status_code == 200  # projects_shared default


def test_export_routes(client: tuple[TestClient, str]) -> None:
    """docx / md / xlsx / zip exports over HTTP — deterministic, labelled DRAFT until approved."""
    import io
    import zipfile

    from docx import Document

    fixtures = Path(__file__).parent / "fixtures" / "change_artifacts"
    c, kb_id = client
    cr_id = _create(c, kb_id)["change_request"]["cr_id"]
    base = f"/change-requests/{cr_id}"
    # Nothing drafted yet → 400, and unknown kinds/formats are rejected.
    assert c.get(f"{base}/artifacts/impact/export?format=docx").status_code == 400
    assert c.get(f"{base}/export.zip").status_code == 400
    assert c.get(f"{base}/artifacts/nope/export").status_code == 404
    assert c.get(f"{base}/artifacts/impact/export?format=pdf").status_code == 422

    # A hand-written impact analysis (a human_edit version, not approved).
    markdown = (fixtures / "BCR-001-impact-analysis.md").read_text(encoding="utf-8")
    put = c.put(f"{base}/artifacts/impact", json={"markdown": markdown, "note": "fixture"})
    assert put.status_code == 200, put.text

    docx = c.get(f"{base}/artifacts/impact/export?format=docx")
    assert docx.status_code == 200
    assert docx.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert 'filename="Impact-Analysis-BCR-001-DRAFT.docx"' in docx.headers["content-disposition"]
    doc = Document(io.BytesIO(docx.content))
    assert doc.paragraphs[0].text == "Impact Analysis"
    assert doc.paragraphs[1].text.endswith("— DRAFT v1 — not approved")

    md = c.get(f"{base}/artifacts/impact/export?format=md")
    assert md.status_code == 200 and md.text.startswith("# Impact Analysis — BCR-001")
    assert 'filename="BCR-001-impact-analysis-DRAFT.md"' in md.headers["content-disposition"]

    xlsx = c.get(f"{base}/artifacts/impact/export?format=xlsx")
    assert xlsx.status_code == 200 and xlsx.content[:2] == b"PK"
    assert 'filename="TC-preview-BCR-001-DRAFT.xlsx"' in xlsx.headers["content-disposition"]
    assert c.get(f"{base}/artifacts/epic/export?format=xlsx").status_code == 400  # not drafted
    stories = (fixtures / "US-008-015-stories.md").read_text(encoding="utf-8")
    assert c.put(f"{base}/artifacts/stories", json={"markdown": stories}).status_code == 200
    stories_zip = c.get(f"{base}/artifacts/stories/export?format=docx")
    assert (
        stories_zip.status_code == 200 and stories_zip.headers["content-type"] == "application/zip"
    )
    with zipfile.ZipFile(io.BytesIO(stories_zip.content)) as zf:
        assert len(zf.namelist()) == 8 and all(n.endswith("-DRAFT.docx") for n in zf.namelist())

    bundle = c.get(f"{base}/export.zip")
    assert bundle.status_code == 200 and bundle.headers["content-type"] == "application/zip"
    with zipfile.ZipFile(io.BytesIO(bundle.content)) as zf:
        names = zf.namelist()
        manifest = zf.read("MANIFEST.txt").decode("utf-8")
    assert (
        "Impact-Analysis-BCR-001-DRAFT.docx" in names and "TC-preview-BCR-001-DRAFT.xlsx" in names
    )
    assert sum(1 for n in names if n.startswith("US-0")) == 8
    assert "markdown/BCR-001-impact-analysis-DRAFT.md" in names
    assert "epic: not drafted — skipped" in manifest and "tdd: not drafted — skipped" in manifest

    # Another signed-in account: shared by default (200), 404 when per-owner isolation is on.
    c.post("/auth/logout")
    c.post(
        "/auth/register",
        json={"email": "other@example.com", "password": "password123", "display_name": "O"},
    )
    assert c.get(f"{base}/export.zip").status_code in (200, 404)


def test_export_filename_header_is_exposed_cross_origin(client: tuple[TestClient, str]) -> None:
    c, _kb_id = client
    response = c.get("/health", headers={"Origin": "http://127.0.0.1:3010"})
    assert (
        "content-disposition" in response.headers.get("access-control-expose-headers", "").lower()
    )
