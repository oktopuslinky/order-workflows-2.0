"""Phase 4 — post-approval change outputs (models, deterministic helpers, engine, export)."""
# ruff: noqa: E501  (long literal file / diagram texts)

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook

from workflow_compiler.agents.change_outputs import ChangeOutputsAgent
from workflow_compiler.change_outputs.code import (
    category_rank,
    check_syntax,
    continue_code,
    extract_code,
    imports_of,
    missing_symbols,
    plan_rewrites,
    resolve_component_file,
    resolve_import,
    signature_summary,
    unified_diff,
)
from workflow_compiler.change_outputs.diagrams import (
    assemble_system_flow,
    balanced,
    check_diagram,
    expected_states,
    mermaid_header,
    plan_diagrams,
    states_in,
)
from workflow_compiler.change_outputs.engine import (
    ChangeOutputsError,
    change_label_of,
    design_summary,
)
from workflow_compiler.change_outputs.export import (
    changes_index,
    export_entries,
    export_zip,
    zip_code_path,
)
from workflow_compiler.change_outputs.models import (
    ChangedFile,
    ChangeOutputs,
    CodeChangeBundle,
    DiagramKind,
    FileChecks,
    FileStatus,
    StageRecord,
    TestCaseDraft,
    TestCaseUpdate,
    TestDocUpdate,
    TestPlanAddendumDraft,
    UpdatedDiagram,
)
from workflow_compiler.change_outputs.tests_doc import (
    export_addendum_docx,
    export_matrix_xlsx,
    merge_test_cases,
    next_tc_ids,
    parse_addendum_meta,
    render_addendum,
)
from workflow_compiler.compiler import ReviewConfig, WorkflowCompiler
from workflow_compiler.docs_export.xlsx_writer import (
    TestCaseRow,
    TestCaseSummary,
    read_test_case_rows,
    write_test_case_matrix,
)
from workflow_compiler.exceptions import ApprovalError
from workflow_compiler.kg import InMemoryKnowledgeBaseStore, KgGrounder, KgService, KnowledgeBase
from workflow_compiler.kg.ingest import zip_folder
from workflow_compiler.llm.providers.mock import MockProvider
from workflow_compiler.models import (
    ChangeSpec,
    ChangeType,
    CompilationProject,
    ComponentChange,
    ComponentKind,
    ProjectStage,
    TemporalActivityDesign,
    TemporalQueryDesign,
    TemporalSignalDesign,
    TemporalWorkflowDesign,
)
from workflow_compiler.project_compiler import ProjectCompiler
from workflow_compiler.storage import InMemoryStateStore
from workflow_compiler.storage.project_store import InMemoryProjectStore

from .test_change_spec import TDD_TEXT

_NO_REVIEW = ReviewConfig(enabled=False)

TYPES_PY = '''"""Shared types."""
from dataclasses import dataclass
from enum import Enum


class OrderStatus(str, Enum):
    RECEIVED = "RECEIVED"
    DISPATCHED = "DISPATCHED"


@dataclass
class ProvisioningResult:
    reservation_id: str
'''
ACTIVITIES_PY = '''"""Activities."""
from src.shared.types import ProvisioningResult


async def provision_order(order_id: str) -> ProvisioningResult:
    return ProvisioningResult(reservation_id=f"RSV-{order_id}")
'''
WORKFLOW_PY = '''"""Workflow."""
from src.activities.order_activities import provision_order
from src.shared.types import OrderStatus


class OrderWorkflow:
    async def run(self, order_id: str) -> OrderStatus:
        await provision_order(order_id)
        return OrderStatus.DISPATCHED

    def get_status(self) -> OrderStatus:
        return OrderStatus.RECEIVED
'''
WORKER_PY = '''"""Worker."""
from src.activities.order_activities import provision_order
from src.workflows.order_workflow import OrderWorkflow

ACTIVITIES = [provision_order]
WORKFLOWS = [OrderWorkflow]
'''
STARTER_PY = '''"""Starter."""
from src.worker import ACTIVITIES
from src.workflows.order_workflow import OrderWorkflow

print(OrderWorkflow, ACTIVITIES)
'''
TESTS_PY = '''"""Tests."""
from src.shared.types import OrderStatus
from src.workflows.order_workflow import OrderWorkflow


def test_status() -> None:
    assert OrderWorkflow().get_status() is OrderStatus.RECEIVED
'''
UTIL_PY = '''"""Unrelated helper."""


def helper() -> int:
    return 1
'''

CORPUS_TEXTS = {
    "existing_Codebase/shared/types.py": TYPES_PY,
    "existing_Codebase/activities/order_activities.py": ACTIVITIES_PY,
    "existing_Codebase/workflows/order_workflow.py": WORKFLOW_PY,
    "existing_Codebase/worker.py": WORKER_PY,
    "existing_Codebase/starter.py": STARTER_PY,
    "existing_Codebase/__init__.py": "",
    "existing_Codebase/activities/__init__.py": "",
    "existing_Codebase/util.py": UTIL_PY,
    "tests/test_order_workflow.py": TESTS_PY,
}

STATE_MMD = """stateDiagram-v2
    [*] --> RECEIVED: capture_order
    RECEIVED --> PROVISIONING: begin provisioning
    PROVISIONING --> DISPATCHED: dispatched
    DISPATCHED --> [*]
"""

FLOW_MD = """# System & Process Flow Diagrams — Orders

Source `.mmd` files live in `docs/diagrams/mermaid/`.

## 1. Order State Machine

Governs every valid transition. See TDD §4.1.

```mermaid
stateDiagram-v2
    [*] --> RECEIVED
```

## 2. End-to-End Sequence (Happy Path)

Shows the calls. See TDD §4.7.

```mermaid
sequenceDiagram
    A->>B: hi
```
"""


def _spec() -> ChangeSpec:
    return ChangeSpec(
        components=[
            ComponentChange(
                name="OrderStatus", kind=ComponentKind.TYPE,
                path="fn:existing_Codebase/shared/types.py:OrderStatus",
                existing="RECEIVED, DISPATCHED", proposed="Add PARTIALLY_PROVISIONED.",
            ),
            ComponentChange(
                name="provision_group", kind=ComponentKind.ACTIVITY, path="",
                proposed="Provision one shipment group.", change_type=ChangeType.ADD,
            ),
            ComponentChange(
                name="get_status", kind=ComponentKind.QUERY,
                path="fn:existing_CodeBase/workflows/order_workflow.py:get_status",
                proposed="Return per-group status.",
            ),
            ComponentChange(
                name="TC-06", kind=ComponentKind.TEST, path="TC-06",
                proposed="Verify per group.",
            ),
            ComponentChange(
                name="order-state-machine.mmd", kind=ComponentKind.DIAGRAM,
                path="doc:Business_Docs/diagrams/mermaid/order-state-machine.mmd",
                proposed="Add PARTIALLY_PROVISIONED and PARTIALLY_DISPATCHED.",
            ),
            ComponentChange(
                name="order-state-machine-partial-shipment.mmd", kind=ComponentKind.DIAGRAM,
                path="", proposed="Group sub-state machine.", change_type=ChangeType.ADD,
            ),
        ]
    )


# ------------------------------------------------------------------ models


def test_change_outputs_round_trip() -> None:
    outputs = ChangeOutputs(
        diagrams=[UpdatedDiagram(name="a.mmd", kind=DiagramKind.STATE, original="x", updated="y")],
        code=CodeChangeBundle(
            files=[ChangedFile(path="p.py", status=FileStatus.MODIFIED, original="a", updated="b",
                               unified_diff="d", checks=FileChecks(ast_ok=True, ruff_ok=None))],
            order=["p.py"], import_root="src", code_root="existing_Codebase",
        ),
        tests_doc=TestDocUpdate(test_cases=[TestCaseRow(tc_id="TC-18", title="t")], new_ids=["TC-18"]),
        stages={"diagrams": StageRecord(status="done", seconds=1.5)},
        timings={"diagrams": 1.5},
    )
    again = ChangeOutputs.model_validate_json(outputs.model_dump_json())
    assert again == outputs
    assert again.done_stages() == ["diagrams"]
    assert again.code.changed()[0].path == "p.py"


# ------------------------------------------------------------------ code helpers


def test_plan_rewrites_selects_dependents_and_orders() -> None:
    plan = plan_rewrites(_spec(), CORPUS_TEXTS)
    assert plan.import_root == "src"
    assert plan.code_root == "existing_Codebase"
    # named: types (OrderStatus), activities (provision_group falls back), workflow (get_status)
    assert plan.order[:3] == [
        "existing_Codebase/shared/types.py",
        "existing_Codebase/activities/order_activities.py",
        "existing_Codebase/workflows/order_workflow.py",
    ]
    # dependents follow: worker (imports activities+workflow) → starter (imports worker) → tests
    assert plan.order[3:] == [
        "existing_Codebase/worker.py",
        "existing_Codebase/starter.py",
        "tests/test_order_workflow.py",
    ]
    assert "existing_Codebase/util.py" in plan.unchanged
    assert "existing_Codebase/__init__.py" in plan.unchanged
    # a new activity never lands in a package __init__ (the shortest "activit…" path)
    assert "existing_Codebase/activities/__init__.py" in plan.unchanged
    assert plan.reasons["existing_Codebase/worker.py"][0].startswith("imports ")
    assert [c.name for c in plan.components_by_file["existing_Codebase/shared/types.py"]] == [
        "OrderStatus"
    ]
    assert plan.imports["existing_Codebase/starter.py"] == [
        "existing_Codebase/worker.py", "existing_Codebase/workflows/order_workflow.py",
    ]


def test_resolve_component_file_and_imports() -> None:
    files = list(CORPUS_TEXTS)
    comp = ComponentChange(name="get_status", kind=ComponentKind.QUERY,
                           path="fn:existing_CodeBase/workflows/order_workflow.py:get_status")
    assert resolve_component_file(comp, files) == "existing_Codebase/workflows/order_workflow.py"
    assert resolve_component_file(ComponentChange(name="TC-06", kind=ComponentKind.TEST), files) is None
    assert resolve_component_file(
        ComponentChange(name="tests/test_order_workflow.py", kind=ComponentKind.TEST), files
    ) == "tests/test_order_workflow.py"
    assert resolve_import("src.shared.types", files) == "existing_Codebase/shared/types.py"
    assert resolve_import("src.worker", files) == "existing_Codebase/worker.py"
    assert resolve_import("temporalio.client", files) is None
    assert "src.worker" in imports_of(STARTER_PY)
    assert category_rank("tests/test_x.py") == 5 and category_rank("a/shared/types.py") == 0


def test_extract_and_continue_code() -> None:
    fenced = extract_code("Sure!\n```python\nx = 1\n```\nDone.")
    assert fenced.found and fenced.closed and fenced.code == "x = 1\n"
    unclosed = extract_code("```python\nx = 1\ny = 2")
    assert unclosed.found and not unclosed.closed and unclosed.code == "x = 1\ny = 2\n"
    merged = continue_code(unclosed.code, "```python\ny = 2\nz = 3\n```")
    assert merged.closed and merged.code == "x = 1\ny = 2\nz = 3\n"
    bare = extract_code("import os\nprint(os.sep)\n")
    assert bare.found and bare.closed
    assert not extract_code("I cannot do that.").found
    ok, err = check_syntax("def f(:\n  pass")
    assert not ok and "line 1" in err
    diff = unified_diff("a.py", "x = 1\n", "x = 2\n")
    assert diff.startswith("--- a/a.py\n+++ b/a.py\n") and "-x = 1" in diff and "+x = 2" in diff
    assert unified_diff("a.py", "same", "same") == ""
    summary = signature_summary(WORKFLOW_PY)
    assert "class OrderWorkflow():" in summary
    assert "async def run(self, order_id: str) -> OrderStatus" in summary
    assert missing_symbols("def provision_group(): ...", ["provision_group", "OrderStatus"]) == [
        "OrderStatus"
    ]


# ------------------------------------------------------------------ diagram helpers


def test_diagram_checks_and_flow_assembly() -> None:
    assert mermaid_header(STATE_MMD) == "statediagram-v2"
    assert mermaid_header("hello\nworld") is None
    assert states_in(STATE_MMD) == {"RECEIVED", "PROVISIONING", "DISPATCHED"}
    assert balanced("flowchart LR\n subgraph A\n  X\n") == ["unbalanced subgraph/end (1 subgraph, 0 end)"]
    required = expected_states(_spec(), states_in(STATE_MMD))
    assert {"RECEIVED", "PARTIALLY_PROVISIONED", "PARTIALLY_DISPATCHED"} <= required
    bad = UpdatedDiagram(name="order-state-machine.mmd", kind=DiagramKind.STATE, updated=STATE_MMD)
    failures = check_diagram(bad, required_states=required)
    assert failures and failures[0].startswith("missing state(s): PARTIALLY_DISPATCHED")
    good = UpdatedDiagram(
        name="x.mmd", kind=DiagramKind.STATE,
        updated=STATE_MMD + "    PROVISIONING --> PARTIALLY_PROVISIONED\n"
        "    PARTIALLY_PROVISIONED --> PARTIALLY_DISPATCHED\n    PARTIALLY_DISPATCHED --> DISPATCHED\n",
    )
    assert check_diagram(good, required_states=required) == []
    requests = plan_diagrams(_spec(), ["Business_Docs/diagrams/mermaid/order-state-machine.mmd",
                                       "Business_Docs/diagrams/mermaid/order-sequence.mmd"])
    assert [r.name for r in requests] == [
        "order-sequence.mmd", "order-state-machine.mmd", "order-state-machine-partial-shipment.mmd",
    ]
    assert requests[2].kind is DiagramKind.STATE_PARTIAL and requests[2].source_path == ""
    diagrams = [
        UpdatedDiagram(name="order-state-machine.mmd", kind=DiagramKind.STATE, updated="stateDiagram-v2\n    A --> B\n"),
        UpdatedDiagram(name="order-sequence.mmd", kind=DiagramKind.SEQUENCE, updated="sequenceDiagram\n    A->>B: x\n"),
        UpdatedDiagram(name="order-state-machine-partial-shipment.mmd", kind=DiagramKind.STATE_PARTIAL,
                       updated="stateDiagram-v2\n    G --> H\n", notes="Group lifecycle."),
    ]
    flow = assemble_system_flow(FLOW_MD, diagrams, {"orderworkflow": "flowchart TD\n  S --> E\n"},
                                change_title="BCR-001")
    assert flow.startswith("# System & Process Flow Diagrams — Orders\n\n_Updated for BCR-001._")
    assert "## 1. Order State Machine" in flow and "    A --> B" in flow
    assert "## 2. End-to-End Sequence (Happy Path)" in flow
    assert "## 3. Workflow Specification Diagram" in flow and "S --> E" in flow
    assert "## 4. Order State Machine Partial Shipment" in flow and "Group lifecycle." in flow
    assert flow.count("```mermaid") == 4


# ------------------------------------------------------------------ test docs


def test_tc_ids_merge_and_addendum_round_trip(tmp_path: Path) -> None:
    assert next_tc_ids(["TC-01", "TC-17"], 2) == ["TC-18", "TC-19"]
    assert next_tc_ids([], 1, start_hint="TC-18") == ["TC-18"]
    existing = [
        TestCaseRow(tc_id="TC-05", title="Provisioning fails", type="Functional", automated="Yes",
                    notes="orig"),
        TestCaseRow(tc_id="TC-06", title="Dispatch fails", type="Functional / Compensation",
                    automated="Yes"),
    ]
    rows, changed, new = merge_test_cases(
        existing,
        [TestCaseDraft(title="Split — two groups", type="functional", automated="yes",
                       linked="US-008", steps="1. a\n2. b", expected="PARTIALLY_PROVISIONED"),
         TestCaseDraft(title="split — two groups", type="Functional")],  # duplicate title
        [TestCaseUpdate(tc_id="TC-06", notes="per group"), TestCaseUpdate(tc_id="TC-99", notes="x")],
        start_hint="TC-18", change_note="BCR-001: new scenario",
    )
    assert changed == ["TC-06"] and new == ["TC-18"]
    assert rows[1].notes == "per group" and rows[0].notes == "orig"
    assert rows[2].tc_id == "TC-18" and rows[2].type == "Functional" and rows[2].automated == "Yes"
    assert rows[2].notes.endswith("BCR-001: new scenario")
    update = TestDocUpdate(
        test_cases=rows, changed_ids=changed, new_ids=new, test_plan_id="TP-ORD-001",
        change_request_id="BCR-001", linked_tdd="TDD-ORD-002", linked_epic="EPIC-002",
        matrix_source="Business_Docs/test-cases/TC-order-workflow.xlsx",
    )
    update.test_plan_addendum_md = render_addendum(
        TestPlanAddendumDraft(out_of_scope_removed=["Partial shipment (BCR-001)"],
                              test_types_added=["Fan-out — per group"]),
        test_plan_id="TP-ORD-001", change_request_id="BCR-001", change_title="Partial shipments",
        linked_tdd="TDD-ORD-002", linked_epic="EPIC-002", changed_ids=changed, new_ids=new, rows=rows,
    )
    md = update.test_plan_addendum_md
    assert md.startswith("# TP-ORD-001 — Addendum — BCR-001 — Partial shipments\n")
    assert parse_addendum_meta(md)["Linked TDD"] == "TDD-ORD-002"
    assert "| TC-18 | Split — two groups | Functional | Yes | US-008 |" in md
    assert "| TC-06 | Dispatch fails | per group |" in md
    # xlsx round trip through the Phase 2 reader
    back = read_test_case_rows(export_matrix_xlsx(update, label="v1"))
    assert [r.tc_id for r in back] == ["TC-05", "TC-06", "TC-18"]
    assert back[2].expected == "PARTIALLY_PROVISIONED"
    # addendum docx opens with the reference-style title and headings
    doc = Document(io.BytesIO(export_addendum_docx(update)))
    texts = [p.text for p in doc.paragraphs]
    assert texts[0] == "Test Plan — Addendum"
    assert any(t == "1. Purpose" for t in texts) and any("Out of Scope" in t for t in texts)
    assert any(t.startswith("Amends:") for t in texts)


# ------------------------------------------------------------------ engine (mock provider)


def _write_corpus(root: Path) -> Path:
    for path, text in CORPUS_TEXTS.items():
        target = root / path
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(text, encoding="utf-8")
    mmd = root / "Business_Docs" / "diagrams" / "mermaid"
    mmd.mkdir(parents=True)
    (mmd / "order-state-machine.mmd").write_text(STATE_MMD, encoding="utf-8")
    (mmd / "order-sequence.mmd").write_text("sequenceDiagram\n    A->>B: hi\n", encoding="utf-8")
    (root / "Business_Docs" / "diagrams" / "system-flow-diagram.md").write_text(FLOW_MD, encoding="utf-8")
    tc = root / "Business_Docs" / "test-cases"
    tc.mkdir(parents=True)
    rows = [TestCaseRow(tc_id=f"TC-{i:02d}", title=f"Case {i}", type="Functional", automated="Yes")
            for i in range(1, 18)]
    (tc / "TC-order-workflow.xlsx").write_bytes(
        write_test_case_matrix(rows, TestCaseSummary(title="Matrix", linked_tdd="TDD-ORD-001"))
    )
    doc = Document()
    doc.add_paragraph("Test Plan")
    doc.add_paragraph("Document ID: TP-ORD-001")
    doc.add_paragraph("3.2 Out of Scope: partial shipment until BCR-001 is implemented.")
    doc.save(tc / "TP-order-workflow-test-plan.docx")
    (root / "docs").mkdir()
    (root / "docs" / "TDD-order.md").write_text("# TDD-ORD-001\n\nOrderWorkflow design.\n", encoding="utf-8")
    return root


@pytest.fixture
def kg(tmp_path: Path) -> KgService:
    return KgService(InMemoryKnowledgeBaseStore(tmp_path / "state"))


@pytest.fixture
async def kb(kg: KgService, tmp_path: Path) -> KnowledgeBase:
    corpus = _write_corpus(tmp_path / "corpus")
    created = await kg.create_from_zip("orders", zip_folder(corpus), owner_id="u1", filename="c.zip")
    return await kg.index(created.kb_id)


def _fence(code: str) -> str:
    return f"```python\n{code}```\n"


NEW_TYPES = TYPES_PY.replace('    DISPATCHED = "DISPATCHED"\n',
                             '    DISPATCHED = "DISPATCHED"\n    PARTIALLY_PROVISIONED = "PARTIALLY_PROVISIONED"\n')
NEW_ACTIVITIES = ACTIVITIES_PY + '''

async def provision_group(order_id: str, group_id: str) -> ProvisioningResult:
    return ProvisioningResult(reservation_id=f"RSV-{order_id}-{group_id}")
'''
BROKEN_WORKFLOW = WORKFLOW_PY.replace("    def get_status(self) -> OrderStatus:", "    def get_status(self -> OrderStatus:")
NEW_WORKFLOW = WORKFLOW_PY.replace("return OrderStatus.RECEIVED", "return OrderStatus.PARTIALLY_PROVISIONED")


def _completions() -> list[str]:
    """One answer per file in rewrite order (+ a repair for the broken workflow)."""
    return [
        _fence(NEW_TYPES),
        # activities: truncated, then continued
        "```python\n" + NEW_ACTIVITIES.split("\n\nasync def provision_group")[0] + "\n",
        _fence("\nasync def provision_group" + NEW_ACTIVITIES.split("\n\nasync def provision_group")[1]),
        _fence(BROKEN_WORKFLOW),  # syntax error → repair round
        _fence(NEW_WORKFLOW),  # the repair
        _fence(WORKER_PY),  # returned verbatim → unchanged
        "no code, sorry",  # starter: no code → unchanged, warning
        _fence(TESTS_PY.replace("OrderStatus.RECEIVED", "OrderStatus.PARTIALLY_PROVISIONED")),
    ]


def _project_compiler(kg: KgService, provider: MockProvider) -> ProjectCompiler:
    inner = WorkflowCompiler(llm_provider=provider, state_store=InMemoryStateStore(), review=_NO_REVIEW)
    return ProjectCompiler(
        llm_provider=provider, workflow_compiler=inner, project_store=InMemoryProjectStore(),
        segmentation_review=False, kg_service=kg,
    )


async def _approved_project(kg: KgService, kb: KnowledgeBase, provider: MockProvider) -> tuple[ProjectCompiler, str]:
    compiler = _project_compiler(kg, provider)
    grounder = KgGrounder(kg, kb.kb_id, kb_name="Orders KB")
    project = await compiler.compile_document(TDD_TEXT, grounder=grounder)
    project.change_spec = _spec()
    project.grounding.change_request_title = "BCR-001 — Partial shipments"  # type: ignore[union-attr]
    await compiler.save_project(project)
    approved = await compiler.approve_spec(project.project_id, accept_incomplete=True)
    assert approved.workflow_ids
    return compiler, project.project_id


async def test_engine_runs_all_stages_with_mock(kg: KgService, kb: KnowledgeBase) -> None:
    provider = MockProvider(script_defaults=True, completions=_completions())
    compiler, pid = await _approved_project(kg, kb, provider)
    saves: list[str] = []
    original_save = compiler._projects.save

    async def spy(project):  # type: ignore[no-untyped-def]
        saves.append(",".join(project.change_outputs.done_stages()) if project.change_outputs else "")
        await original_save(project)

    compiler._projects.save = spy  # type: ignore[method-assign]
    project = await compiler.generate_change_outputs(pid)
    outputs = project.change_outputs
    assert outputs is not None
    assert outputs.done_stages() == ["diagrams", "code", "tests_doc"]
    assert set(outputs.timings) == {"diagrams", "code", "tests_doc"}
    assert "change_outputs" in project.stage_timings
    # persisted after each stage (and after every file): more saves than stages
    assert len(saves) > 3 and saves[-1] == "diagrams,code,tests_doc"

    # diagrams: originals regenerated + companion + the spec's workflow diagram, flow doc assembled
    by_name = {d.name: d for d in outputs.diagrams}
    assert {"order-state-machine.mmd", "order-sequence.mmd",
            "order-state-machine-partial-shipment.mmd"} <= set(by_name)
    assert by_name["order-state-machine.mmd"].original == STATE_MMD
    assert "PARTIALLY_PROVISIONED" in by_name["order-state-machine.mmd"].updated
    assert by_name["order-state-machine-partial-shipment.mmd"].original is None
    # the mock's sequence diagram was not returned → original kept, check recorded
    assert by_name["order-sequence.mmd"].updated == "sequenceDiagram\n    A->>B: hi\n"
    assert by_name["order-sequence.mmd"].checks == ["model returned no diagram"]
    assert any(d.kind is DiagramKind.WORKFLOW for d in outputs.diagrams)
    assert "## 1. Order State Machine" in outputs.system_flow_md
    assert "Workflow Specification Diagram" in outputs.system_flow_md
    assert "Business_Docs/diagrams/mermaid/order-state-machine.mmd" in outputs.provenance

    # code: ordered rewrite, continuation, repair, unchanged copies, diffs
    code = outputs.code
    assert code.order == [
        "existing_Codebase/shared/types.py",
        "existing_Codebase/activities/order_activities.py",
        "existing_Codebase/workflows/order_workflow.py",
        "existing_Codebase/worker.py",
        "existing_Codebase/starter.py",
        "tests/test_order_workflow.py",
    ]
    files = {f.path: f for f in code.files}
    assert files["existing_Codebase/shared/types.py"].status is FileStatus.MODIFIED
    assert files["existing_Codebase/shared/types.py"].updated == NEW_TYPES
    assert "+    PARTIALLY_PROVISIONED" in files["existing_Codebase/shared/types.py"].unified_diff
    acts = files["existing_Codebase/activities/order_activities.py"]
    assert acts.checks.truncated and acts.checks.ast_ok and "provision_group" in acts.updated
    wf = files["existing_Codebase/workflows/order_workflow.py"]
    assert wf.checks.repaired and wf.checks.ast_ok and wf.updated == NEW_WORKFLOW
    assert files["existing_Codebase/worker.py"].status is FileStatus.UNCHANGED
    assert files["existing_Codebase/starter.py"].status is FileStatus.UNCHANGED
    assert any("starter.py: model returned no code" in w for w in outputs.warnings)
    assert files["existing_Codebase/util.py"].status is FileStatus.UNCHANGED
    assert files["existing_Codebase/util.py"].reason == ""
    assert files["tests/test_order_workflow.py"].status is FileStatus.MODIFIED
    # later files saw the earlier outputs' signatures
    prompts = [p for kind, p in provider.calls if kind == "complete"]
    workflow_prompt = next(p for p in prompts if "FILE TO REWRITE: `existing_Codebase/workflows/order_workflow.py`" in p)
    assert "provision_group" in workflow_prompt and "PARTIALLY_PROVISIONED" in workflow_prompt
    assert "SIGNATURES OF FILES ALREADY REWRITTEN" in workflow_prompt

    # tests_doc: ids from the matrix, updated TC-06, addendum
    tests = outputs.tests_doc
    assert tests.new_ids == ["TC-18"] and tests.changed_ids == ["TC-06"]
    assert len(tests.test_cases) == 18 and tests.test_cases[-1].tc_id == "TC-18"
    assert tests.test_plan_id == "TP-ORD-001" and tests.change_request_id == "BCR-001"
    assert tests.test_plan_addendum_md.startswith("# TP-ORD-001 — Addendum — BCR-001")
    assert tests.matrix_source == "Business_Docs/test-cases/TC-order-workflow.xlsx"

    # export zip: README layout + docs + CHANGES.md
    data = export_zip(outputs, project_id=pid, label="BCR-001")
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        names = archive.namelist()
        assert "src/shared/types.py" in names and "src/util.py" in names
        assert "tests/test_order_workflow.py" in names
        assert "docs/diagrams/mermaid/order-state-machine.mmd" in names
        assert "docs/diagrams/mermaid/order-state-machine-partial-shipment.mmd" in names
        assert "docs/diagrams/system-flow-diagram.md" in names
        assert "docs/test-cases/TC-order-workflow.xlsx" in names
        assert "docs/test-cases/TP-ORD-001-addendum-BCR-001.docx" in names
        assert "docs/test-cases/TP-ORD-001-addendum-BCR-001.md" in names
        assert "changes.patch" in names and "CHANGES.md" in names
        assert archive.read("src/shared/types.py").decode() == NEW_TYPES
        index = archive.read("CHANGES.md").decode()
        assert "| existing_Codebase/shared/types.py | src/shared/types.py | modified |" in index
        assert "TC-18" in index
    assert export_zip(outputs, project_id=pid, label="BCR-001") == data  # byte-stable
    assert zip_code_path("existing_Codebase/a/b.py", code_root="existing_Codebase", import_root="src") == "src/a/b.py"
    entries = export_entries(outputs, project_id=pid)
    assert entries[-1].path == "CHANGES.md"
    assert "## Stages" in changes_index(outputs, project_id=pid)


async def test_engine_stage_failure_keeps_other_stages(kg: KgService, kb: KnowledgeBase) -> None:
    provider = MockProvider(script_defaults=True)  # no completions → files unchanged, fine
    compiler, pid = await _approved_project(kg, kb, provider)
    project = await compiler.generate_change_outputs(pid, stages=["diagrams"])
    assert project.change_outputs is not None
    assert project.change_outputs.done_stages() == ["diagrams"]
    # a second run of another stage keeps the first
    project = await compiler.generate_change_outputs(pid, stages=["tests_doc"])
    assert project.change_outputs is not None
    assert project.change_outputs.done_stages() == ["diagrams", "tests_doc"]
    # a stage whose model call fails is recorded failed; the run raises after finishing the rest
    strict = MockProvider(script_defaults=False)
    compiler2 = _project_compiler(kg, strict)
    compiler2._projects = compiler._projects
    with pytest.raises(ChangeOutputsError, match="diagrams"):
        await compiler2.generate_change_outputs(pid, stages=["diagrams", "code"])
    reloaded = await compiler.load_project(pid)
    assert reloaded.change_outputs is not None
    assert reloaded.change_outputs.stages["diagrams"].status == "failed"
    assert reloaded.change_outputs.stages["code"].status == "done"  # no code from the model = unchanged
    assert reloaded.change_outputs.stages["tests_doc"].status == "done"
    with pytest.raises(ApprovalError, match="not grounded"):
        plain = await compiler.compile_document("A plain document about invoices.")
        await compiler.generate_change_outputs(plain.project_id)


def test_design_summary_and_change_label() -> None:
    design = TemporalWorkflowDesign(
        workflow_name="OrderWorkflow", task_queue="orders",
        activities=[TemporalActivityDesign(name="provision_group", result_type="ProvisioningResult")],
        signals=[TemporalSignalDesign(name="cancel_shipment_group", payload=["group_id"])],
        queries=[TemporalQueryDesign(name="get_status", returns="OrderState")],
    )
    text = design_summary([design])
    assert "Workflow OrderWorkflow (task queue orders)" in text
    assert "activity provision_group() -> ProvisioningResult" in text
    assert "signal cancel_shipment_group(group_id)" in text and "query get_status() -> OrderState" in text
    from workflow_compiler.models import ProjectGrounding

    project = CompilationProject(document_text="Design for BCR-001 partial shipments.",
                                 grounding=ProjectGrounding(change_request_title="Partial"))
    assert change_label_of(project) == "BCR-001"
    project.document_text = "nothing"
    project.change_request_id = "abcdef0123"
    assert change_label_of(project) == "abcdef01"


async def test_agent_requires_llm() -> None:
    from workflow_compiler.exceptions import CompilationError

    with pytest.raises(CompilationError):
        await ChangeOutputsAgent(None).repair_file(path="a.py", code="x", error="e")


def test_stage_enum_and_project_stage_import() -> None:
    assert ProjectStage.COMPLETED.value == "completed"
    wb = Workbook()
    assert wb.active is not None


def test_auto_import_known_names() -> None:
    from workflow_compiler.change_outputs.code import auto_import, undefined_names

    ruff_out = "F821 Undefined name `timedelta`\n  --> x.py:3\nF821 Undefined name `timedelta`\nF821 Undefined name `RetryPolicy`\nF821 Undefined name `mystery`"
    assert undefined_names(ruff_out) == ["timedelta", "RetryPolicy", "mystery"]
    code = '"""Doc."""\nfrom __future__ import annotations\n\nimport logging\n\nX = timedelta(seconds=1)\n'
    fixed, added = auto_import(code, undefined_names(ruff_out))
    assert added == ["from datetime import timedelta", "from temporalio.common import RetryPolicy"]
    assert fixed.split("\n")[4] == "from datetime import timedelta"  # after the last import
    assert "mystery" not in " ".join(added)
    # no imports at all: after the docstring
    fixed2, added2 = auto_import('"""Doc."""\nX = uuid.uuid4()\n', ["uuid"])
    assert fixed2.startswith('"""Doc."""\nimport uuid\n') and added2 == ["import uuid"]
    # already present: nothing added
    assert auto_import("import uuid\nX = uuid.uuid4()\n", ["uuid"])[1] == []


def test_missing_imports_against_rewritten_siblings() -> None:
    from workflow_compiler.change_outputs.code import exported_names, missing_imports

    siblings = {"existing_Codebase/activities/order_activities.py": NEW_ACTIVITIES}
    assert {"provision_order", "provision_group", "ProvisioningResult"} <= exported_names(NEW_ACTIVITIES)
    code = "from src.activities.order_activities import provision_group, dispatch_group\n"
    assert missing_imports(code, siblings, list(CORPUS_TEXTS)) == {
        "existing_Codebase/activities/order_activities.py": ["dispatch_group"]
    }
    # imports from files that were not rewritten (or third parties) are not checked
    assert missing_imports("from src.shared.types import Nope\nimport temporalio\n", siblings, list(CORPUS_TEXTS)) == {}


def test_dataclass_problems() -> None:
    from workflow_compiler.change_outputs.code import dataclass_problems

    bad = (
        "from dataclasses import dataclass, field\n\n@dataclass\nclass D:\n    a: str\n"
        "    b: str = field(default='')\n    c: str\n    a: str = ''\n"
    )
    problems = dataclass_problems(bad)
    assert any("non-default field 'c'" in p for p in problems)
    assert any("'a' is declared twice" in p for p in problems)
    good = "from dataclasses import dataclass, field\n\n@dataclass\nclass D:\n    a: str\n    b: list = field(default_factory=list)\n"
    assert dataclass_problems(good) == []


def test_llm_plans_tolerate_stray_items() -> None:
    from workflow_compiler.change_outputs.models import DiagramUpdatePlan, TestCaseUpdatePlan

    plan = DiagramUpdatePlan.model_validate(
        {"diagrams": [{"name": "a.mmd", "mermaid": "stateDiagram-v2"}, "notes", "Added X"], "notes": "n"}
    )
    assert [d.name for d in plan.diagrams] == ["a.mmd"]
    tc = TestCaseUpdatePlan.model_validate({"new_cases": ["x", {"title": "t"}], "updated_cases": "nope", "addendum": "text"})
    assert [c.title for c in tc.new_cases] == ["t"] and tc.updated_cases == [] and tc.addendum.notes == []


def test_corpus_exports_feed_auto_import() -> None:
    from workflow_compiler.change_outputs.code import auto_import, corpus_exports

    exports = corpus_exports(CORPUS_TEXTS, code_root="existing_Codebase", import_root="src")
    assert exports["ProvisioningResult"] == "from src.shared.types import ProvisioningResult"
    assert exports["provision_order"] == "from src.activities.order_activities import provision_order"
    assert "test_status" not in exports  # tests are not package modules
    fixed, added = auto_import("x = ProvisioningResult('r')\n", ["ProvisioningResult"], exports)
    assert added == ["from src.shared.types import ProvisioningResult"] and fixed.startswith("from src.shared")


# --------------------------------------------------------------------------- Phase 5 additions


def test_normalise_style_generics_and_blank_lines() -> None:
    from workflow_compiler.change_outputs.code import normalise_style

    original = (
        '"""Types."""\nfrom dataclasses import dataclass\n\n\n@dataclass\nclass A:\n'
        "    items: list[str]\n    note: str | None = None\n\n\ndef f() -> list[int]:\n    return []\n"
    )
    updated = (
        '"""Types."""\nfrom dataclasses import dataclass\nfrom typing import List, Optional\n'
        "@dataclass\nclass A:\n    items: List[str]\n    note: Optional[str] = None\n"
        "def f() -> List[int]:\n    return []   \n\n\n\n\n"
    )
    text, changed = normalise_style(original, updated)
    assert changed
    assert "from typing import" not in text
    assert "items: list[str]" in text and "note: str | None = None" in text
    assert "def f() -> list[int]:" in text
    assert "\n\n\n@dataclass\nclass A:" in text and "\n\n\ndef f()" in text
    assert text.endswith("return []\n") and "\n\n\n\n" not in text
    # An original that itself uses typing generics is left alone (rule not in force).
    legacy = "from typing import List\n\n\ndef f() -> List[int]:\n    return []\n"
    same, changed2 = normalise_style(legacy, legacy)
    assert not changed2 and same == legacy
    # Decorators stay glued to their definitions.
    stacked = "import x\n@a\n@b\ndef g():\n    pass\n"
    fixed, _ = normalise_style("import x\n\n\n@a\ndef g():\n    pass\n", stacked)
    assert "\n\n\n@a\n@b\ndef g():" in fixed


async def test_smoke_reports_import_errors(tmp_path: Path) -> None:
    from workflow_compiler.change_outputs.smoke import bundle_layout, module_names, run_smoke

    bundle = CodeChangeBundle(
        order=["existing_Codebase/shared/types.py", "existing_Codebase/workflows/order_workflow.py"],
        import_root="src",
        code_root="existing_Codebase",
        files=[
            ChangedFile(path="existing_Codebase/__init__.py", updated=""),
            ChangedFile(path="existing_Codebase/shared/__init__.py", updated=""),
            ChangedFile(path="existing_Codebase/shared/types.py", updated="X = 1\n"),
            ChangedFile(path="existing_Codebase/workflows/__init__.py", updated=""),
            ChangedFile(
                path="existing_Codebase/workflows/order_workflow.py",
                updated="from src.shared.types import MISSING\n",
            ),
            ChangedFile(path="tests/test_x.py", updated="from src.shared.types import X\n"),
            ChangedFile(path="existing_Codebase/gone.py", status=FileStatus.REMOVED, updated=""),
        ],
    )
    layout = bundle_layout(bundle)
    assert "src/shared/types.py" in layout and "existing_Codebase/gone.py" not in layout
    assert module_names(layout, bundle)[:2] == ["src.shared.types", "src.workflows.order_workflow"]
    result = await run_smoke(bundle)
    assert result.status == "failed"
    assert result.compiled == len(layout) and not result.compile_errors
    assert "src.shared.types" in result.imported and "tests.test_x" in result.imported
    assert "src.workflows.order_workflow" in result.import_errors
    assert "MISSING" in result.import_errors["src.workflows.order_workflow"]
    # A passing bundle.
    for f in bundle.files:
        if f.path.endswith("order_workflow.py"):
            f.updated = "from src.shared.types import X\n"
    assert (await run_smoke(bundle)).status == "passed"
    # A syntax error is a compile error (and an import error).
    bundle.files[2].updated = "X = (\n"
    broken = await run_smoke(bundle)
    assert broken.status == "failed" and broken.compile_errors
    # An interpreter that cannot start → skipped with a note, never an exception.
    skipped = await run_smoke(bundle, python=str(tmp_path / "no-such-python"))
    assert skipped.status == "skipped" and "smoke interpreter" in skipped.note


async def test_engine_second_repair_round_and_smoke(kg: KgService, kb: KnowledgeBase) -> None:
    """A file that is still broken after the first repair gets a second, targeted round."""
    completions = [
        _fence(NEW_TYPES),
        _fence(NEW_ACTIVITIES),
        _fence(BROKEN_WORKFLOW),  # syntax error → repair round 1
        _fence(BROKEN_WORKFLOW),  # still broken → repair round 2
        _fence(NEW_WORKFLOW),  # fixed
        _fence(WORKER_PY),
        _fence(STARTER_PY),
        _fence(TESTS_PY),
    ]
    provider = MockProvider(script_defaults=True, completions=completions)
    compiler, pid = await _approved_project(kg, kb, provider)
    project = await compiler.generate_change_outputs(pid, stages=["code"])
    outputs = project.change_outputs
    assert outputs is not None
    files = {f.path: f for f in outputs.code.files}
    wf = files["existing_Codebase/workflows/order_workflow.py"]
    assert wf.checks.repaired and wf.checks.repair_rounds == 2 and wf.checks.ast_ok
    assert len(wf.checks.problems) == 2
    assert all(p.startswith("SyntaxError") for p in wf.checks.problems)
    assert wf.updated == NEW_WORKFLOW
    # The repair prompt carried the verdict.
    repairs = [
        p for kind, p in provider.calls
        if kind == "complete" and "failed a deterministic check" in p
    ]
    assert len(repairs) == 2 and "SyntaxError" in repairs[0]
    # The bundle smoke ran in a subprocess and was recorded.
    smoke = outputs.code.smoke
    assert smoke is not None and smoke.status in {"passed", "failed"}
    assert smoke.compiled == len([f for f in outputs.code.files if f.path.endswith(".py")])
    assert "src.shared.types" in smoke.modules


async def test_engine_repair_rounds_zero_disables_repair(kg: KgService, kb: KnowledgeBase) -> None:
    completions = [
        _fence(NEW_TYPES), _fence(NEW_ACTIVITIES), _fence(BROKEN_WORKFLOW),
        _fence(WORKER_PY), _fence(STARTER_PY), _fence(TESTS_PY),
    ]
    provider = MockProvider(script_defaults=True, completions=completions)
    inner = WorkflowCompiler(
        llm_provider=provider, state_store=InMemoryStateStore(), review=_NO_REVIEW
    )
    compiler = ProjectCompiler(
        llm_provider=provider, workflow_compiler=inner, project_store=InMemoryProjectStore(),
        segmentation_review=False, kg_service=kg, change_outputs_repair_rounds=0,
        change_outputs_smoke=False,
    )
    grounder = KgGrounder(kg, kb.kb_id, kb_name="Orders KB")
    project = await compiler.compile_document(TDD_TEXT, grounder=grounder)
    project.change_spec = _spec()
    await compiler.save_project(project)
    await compiler.approve_spec(project.project_id, accept_incomplete=True)
    project = await compiler.generate_change_outputs(project.project_id, stages=["code"])
    outputs = project.change_outputs
    assert outputs is not None
    wf = {f.path: f for f in outputs.code.files}["existing_Codebase/workflows/order_workflow.py"]
    assert not wf.checks.repaired and wf.checks.repair_rounds == 0 and not wf.checks.ast_ok
    assert outputs.code.smoke is None
    assert any("does not parse after 0 repair round(s)" in w for w in outputs.warnings)


def test_time_saved_buckets_change_outputs() -> None:
    from workflow_compiler.config import Settings
    from workflow_compiler.metrics import compute_time_saved

    project = CompilationProject(document_text="x", stage_timings={"change_outputs": 1800.0})
    report = compute_time_saved(project, Settings().baseline_hours)
    assert report is not None
    row = next(r for r in report.rows if r.category == "change_outputs")
    assert row.human_baseline_hours == 16.0 and "Change outputs" in row.label
