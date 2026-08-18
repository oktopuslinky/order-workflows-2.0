"""Golden-structure tests for the Word / Excel export of change artifacts.

The produced files are opened with python-docx / openpyxl and checked against
the reference conventions encoded in ``tests/fixtures/change_artifacts/
reference_headings.json`` (from the research digest §5): title/subtitle sizes,
heading texts in order, table headers + ``2F5496`` shading, checklist glyphs,
Consolas inline code, xlsx sheet names / columns / summary totals, and the zip
manifest. Inputs are the Phase 1 fixtures (live Nemotron output, approved).
"""

from __future__ import annotations

import io
import json
import re
import zipfile
from datetime import UTC, datetime
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook

from workflow_compiler.change.parse import parse_epic, parse_impact, parse_stories, parse_tdd
from workflow_compiler.docs_export import (
    DocxWriter,
    TestCaseRow,
    TestCaseSummary,
    export_artifact,
    export_change_request,
    export_epic,
    export_impact,
    export_stories,
    export_story,
    export_tdd,
    export_test_case_preview,
    read_test_case_rows,
    render_markdown,
    write_test_case_matrix,
)
from workflow_compiler.docs_export.artifacts import (
    epic_filename,
    export_label,
    impact_filename,
    preview_test_case_rows,
    slugify,
    story_filename,
    tdd_filename,
)
from workflow_compiler.docs_export.docx_writer import parse_inline
from workflow_compiler.docs_export.markdown_to_docx import markdown_document, parse_blocks
from workflow_compiler.models.change import (
    Artifact,
    ArtifactKind,
    ArtifactStatus,
    ChangeRequest,
    VersionSource,
)

FIXTURES = Path(__file__).parent / "fixtures" / "change_artifacts"
REF = json.loads((FIXTURES / "reference_headings.json").read_text(encoding="utf-8"))
XLSX_REF = (
    Path(__file__).parent.parent
    / "examples/knowledge_bases/order-lifecycle/Business_Docs/test-cases/TC-order-workflow.xlsx"
)


def _md(name: str) -> str:
    return (FIXTURES / name).read_text(encoding="utf-8")


def _doc(data: bytes) -> Document:  # type: ignore[valid-type]
    return Document(io.BytesIO(data))


def _headings(doc: Document) -> list[tuple[str, str]]:  # type: ignore[valid-type]
    return [(p.style.name, p.text) for p in doc.paragraphs if p.style.name.startswith("Heading")]


def _shading(cell) -> str | None:  # type: ignore[no-untyped-def]
    shd = cell._tc.find(".//" + qn("w:shd"))
    return None if shd is None else shd.get(qn("w:fill"))


def _meta_labels(doc: Document) -> list[str]:  # type: ignore[valid-type]
    labels: list[str] = []
    for p in doc.paragraphs:
        if p.runs and p.runs[0].bold and p.runs[0].text.endswith(": "):
            labels.append(p.runs[0].text[:-2])
    return labels


# --------------------------------------------------------------------------- #
# Writer primitives + markdown converter
# --------------------------------------------------------------------------- #


def test_parse_inline_spans() -> None:
    spans = parse_inline("Call `get_status()` on **the** *order* now")
    assert [(s.text, s.code, s.bold, s.italic) for s in spans] == [
        ("Call ", False, False, False),
        ("get_status()", True, False, False),
        (" on ", False, False, False),
        ("the", False, True, False),
        (" ", False, False, False),
        ("order", False, False, True),
        (" now", False, False, False),
    ]
    # Unbalanced markers and list types stay literal.
    assert [s.text for s in parse_inline("list[ProvisioningResult] * 2")] == [
        "list[ProvisioningResult] * 2"
    ]


def test_markdown_blocks_and_render() -> None:
    md = (
        "## Heading\n\nA paragraph with `code` and **bold**.\n\n"
        "- one\n- two\n  continued\n- [ ] open\n- [x] done\n\n"
        "1. first\n2. second\n\n"
        "| A | B |\n| --- | --- |\n| x<br>y | p \\| q |\n\n"
        "```python\nprint('hi')\nx = 1\n```\n\n> a note\n\n**Label:** value\n"
    )
    kinds = [b.kind for b in parse_blocks(md)]
    assert kinds == [
        "heading",
        "paragraph",
        "bullet",
        "bullet",
        "check",
        "check",
        "numbered",
        "numbered",
        "table",
        "code",
        "quote",
        "meta",
    ]
    assert parse_blocks(md)[3].text == "two continued"

    writer = DocxWriter()
    render_markdown(writer, md, heading_offset=1)
    doc = _doc(writer.bytes())
    assert _headings(doc) == [("Heading 2", "Heading")]
    texts = [p.text for p in doc.paragraphs]
    assert "A paragraph with code and bold." in texts
    assert "☐  open" in texts and "☑  done" in texts
    numbered = [p for p in doc.paragraphs if p.text in ("first", "second")]
    assert len(numbered) == 2 and all(p._p.pPr.numPr is not None for p in numbered)
    assert "Label: value" in texts
    bullets = [p for p in doc.paragraphs if p.style.name == "List Paragraph"]
    assert [p.text for p in bullets][:2] == ["one", "two continued"]
    assert bullets[0]._p.pPr.numPr is not None
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["A", "B"]
    assert _shading(table.rows[0].cells[0]) == "2F5496"
    assert table.rows[1].cells[0].text == "x\ny" and table.rows[1].cells[1].text == "p | q"
    code = next(p for p in doc.paragraphs if p.text.startswith("print"))
    assert all(r.font.name == "Consolas" for r in code.runs)
    inline = next(p for p in doc.paragraphs if p.text.startswith("A paragraph"))
    assert [r.font.name for r in inline.runs if r.text == "code"] == ["Consolas"]
    assert [r.bold for r in inline.runs if r.text == "bold"] == [True]


def test_markdown_document_fallback() -> None:
    doc = _doc(markdown_document("# Title — Sub\n\ntext\n\n## Sec\n\n- a\n", doc_type="Doc"))
    assert doc.paragraphs[0].text == "Doc" and doc.paragraphs[0].runs[0].font.size.pt == 22
    assert doc.paragraphs[1].text == "Title — Sub" and doc.paragraphs[1].runs[0].font.size.pt == 14
    assert _headings(doc) == [("Heading 1", "Sec")]


def test_writer_is_deterministic() -> None:
    doc = parse_epic(_md("EPIC-002.md"))
    assert export_epic(doc, label="Approved v2") == export_epic(doc, label="Approved v2")


# --------------------------------------------------------------------------- #
# Artifact documents vs the reference conventions
# --------------------------------------------------------------------------- #


def test_impact_docx_structure() -> None:
    doc_model = parse_impact(_md("BCR-001-impact-analysis.md"))
    doc = _doc(export_impact(doc_model, label="Approved v6 (2026-08-18)"))
    ref = REF["impact"]
    assert doc.paragraphs[0].text == ref["title"]
    run = doc.paragraphs[0].runs[0]
    assert run.bold and run.font.size.pt == REF["common"]["title_pt"]
    assert doc.paragraphs[1].text.startswith("BCR-001 — Partial Shipment Support")
    assert doc.paragraphs[1].runs[0].font.size.pt == REF["common"]["subtitle_pt"]
    assert _meta_labels(doc)[:4] == ref["meta_labels"]
    assert "Export: Approved v6 (2026-08-18)" in [p.text for p in doc.paragraphs]
    assert [text for _style, text in _headings(doc)] == ref["sections"]
    assert all(style == "Heading 1" for style, _ in _headings(doc))
    headers = [[c.text for c in t.rows[0].cells] for t in doc.tables]
    assert ref["tables"]["2. Requirements Assessment"] in headers
    assert ref["tables"]["3. Affected Components"] in headers
    for table in doc.tables:
        assert _shading(table.rows[0].cells[0]) == REF["common"]["table_header_fill"]
        hdr_run = table.rows[0].cells[0].paragraphs[0].runs[0]
        assert (
            hdr_run.bold and str(hdr_run.font.color.rgb) == REF["common"]["table_header_font_color"]
        )
        assert _shading(table.rows[1].cells[0]) == "FFFFFF"
        assert table.rows[0]._tr.trPr.find(qn("w:tblHeader")) is not None
    # KG appendix rows and the Sources footer are present.
    appendix = doc.tables[-1]
    assert [c.text for c in appendix.rows[0].cells] == ["Hops", "Type", "Node", "Path", "Via"]
    assert len(appendix.rows) - 1 == len(doc_model.kg_rows) > 50
    sources = [
        p
        for p in doc.paragraphs
        if p.style.name == "List Paragraph" and p.runs and p.runs[0].font.name == "Consolas"
    ]
    assert len(sources) == len(doc_model.sources) > 5
    assert impact_filename(doc_model) == "Impact-Analysis-BCR-001.docx"
    assert impact_filename(doc_model, approved=False) == "Impact-Analysis-BCR-001-DRAFT.docx"


def test_epic_docx_structure() -> None:
    doc_model = parse_epic(_md("EPIC-002.md"))
    doc = _doc(export_epic(doc_model, label="Approved v2"))
    ref = REF["epic"]
    assert doc.paragraphs[0].text == ref["title"] and doc.paragraphs[0].runs[0].font.size.pt == 22
    assert doc.paragraphs[1].text == doc_model.title
    assert doc.paragraphs[2].text == ""  # blank line before the metadata block
    assert _meta_labels(doc)[:5] == ref["meta_labels"]
    headings = _headings(doc)
    assert [text for _s, text in headings] == [*ref["sections"], "Sources"]
    assert {style for style, _ in headings} == {ref["heading_style"]}
    by_heading: dict[str, list[str]] = {}
    current = ""
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            current = p.text
        else:
            by_heading.setdefault(current, []).append(p.text)
    dod = [t for t in by_heading["Definition of Done"] if t]
    assert dod and all(
        t.startswith((REF["common"]["checked"], REF["common"]["unchecked"])) for t in dod
    )
    assert len(dod) == len(doc_model.dod)
    headers = [[c.text for c in t.rows[0].cells] for t in doc.tables]
    for name, columns in ref["tables"].items():
        assert columns in headers, name
    story_map = doc.tables[headers.index(ref["tables"]["Story Map"])]
    assert [r.cells[0].text for r in story_map.rows[1:]] == [s.id for s in doc_model.story_map]
    assert story_map.rows[1].cells[0].text.startswith("US-0")
    bullets = [p for p in doc.paragraphs if p.style.name == REF["common"]["bullet_style"]]
    assert len(bullets) >= len(doc_model.value) + len(doc_model.capabilities)
    assert epic_filename(doc_model).startswith("EPIC-002-") and epic_filename(doc_model).endswith(
        ".docx"
    )


def test_story_docx_structure() -> None:
    stories = parse_stories(_md("US-008-015-stories.md"))
    assert len(stories.stories) == 8
    ref = REF["story"]
    files = export_stories(stories, label="Approved v1")
    assert len(files) == 8
    for entry, story in zip(files, stories.stories, strict=True):
        assert re.match(ref["filename_pattern"], entry.filename), entry.filename
        assert entry.filename.startswith(f"{story.id}-")
    doc = _doc(files[0].data)
    story = stories.stories[0]
    assert re.match(ref["title_pattern"], doc.paragraphs[0].text)
    assert doc.paragraphs[0].runs[0].font.size.pt == 22 and doc.paragraphs[0].runs[0].bold
    assert doc.paragraphs[1].text == ""  # no subtitle for an approved story
    assert _meta_labels(doc)[:3] == ref["meta_labels"]
    headings = _headings(doc)
    assert [text for _s, text in headings][:3] == ref["sections"]
    assert all(style == ref["heading_style"] for style, _ in headings)
    texts = [p.text for p in doc.paragraphs]
    i = texts.index("Story")
    assert texts[i + 1].startswith("As ") and texts[i + 2].lower().startswith("i want")
    subject = doc.paragraphs[i + 1].runs[1]
    assert subject.bold  # "As **a Fulfilment Operator**"
    acs = [t for t in texts if t.startswith(REF["common"]["unchecked"])]
    assert len(acs) == len(story.acceptance) >= 2
    assert story_filename(story) == f"{story.id}-{slugify(story.title)}.docx"


def test_story_draft_label() -> None:
    stories = parse_stories(_md("US-008-015-stories.md"))
    doc = _doc(export_story(stories.stories[0], label="DRAFT v1 — not approved", approved=False))
    assert doc.paragraphs[1].text == "DRAFT v1 — not approved"
    assert "Export: DRAFT v1 — not approved" in [p.text for p in doc.paragraphs]


def test_tdd_docx_structure() -> None:
    doc_model = parse_tdd(_md("TDD-ORD-002.md"))
    doc = _doc(export_tdd(doc_model, label="Approved v2"))
    ref = REF["tdd"]
    assert doc.paragraphs[0].text == ref["title"] and doc.paragraphs[0].runs[0].font.size.pt == 22
    assert doc.paragraphs[1].text == doc_model.title
    assert doc.paragraphs[1].runs[0].font.size.pt == 14
    labels = _meta_labels(doc)
    assert [lbl for lbl in labels if lbl in ref["meta_labels"]] == ref["meta_labels"]
    numbered = [
        (s, t) for s, t in _headings(doc) if s in ("Heading 1", "Heading 2") and t[0].isdigit()
    ]
    assert numbered == [tuple(pair) for pair in ref["sections"]]
    h3 = [t for s, t in _headings(doc) if s == "Heading 3"]
    assert h3 == ["Existing", "Proposed"] * len(doc_model.sections)
    # Bodies keep their tables / inline code (the 4.2 activities table).
    headers = [[c.text for c in t.rows[0].cells] for t in doc.tables]
    assert any(h[0] == "Activity" for h in headers)
    consolas = [r.text for p in doc.paragraphs for r in p.runs if r.font.name == "Consolas"]
    assert "PARTIALLY_PROVISIONED" in consolas or "PARTIALLY_DISPATCHED" in consolas
    assert _headings(doc)[-1] == ("Heading 1", "Sources")
    assert tdd_filename(doc_model) == "TDD-ORD-002-orderworkflow-temporal-implementation.docx"


# --------------------------------------------------------------------------- #
# xlsx
# --------------------------------------------------------------------------- #


def test_xlsx_matrix_round_trip_and_summary() -> None:
    rows = [
        TestCaseRow(
            tc_id="TC-01", title="Happy path", type="Functional", automated="Yes", linked="US-001"
        ),
        TestCaseRow(
            tc_id="TC-02", title="Manual chaos", type="Reliability", automated="Manual (chaos test)"
        ),
        TestCaseRow(
            tc_id="TC-18", title="New one", type="Functional / Compensation", automated="Planned"
        ),
    ]
    summary = TestCaseSummary(
        title="Test Case Matrix — X",
        linked_tdd="TDD-ORD-002",
        linked_epic="EPIC-002",
        automation="tests/",
        notes=["n1", "n2"],
    )
    data = write_test_case_matrix(rows, summary)
    wb = load_workbook(io.BytesIO(data))
    ref = REF["test_cases_xlsx"]
    assert wb.sheetnames == ref["sheets"]
    ws = wb["Test Cases"]
    assert [c.value for c in ws[1]] == ref["columns"]
    assert (
        ws["A1"].fill.fgColor.rgb.endswith("2F5496")
        and ws["A1"].font.bold
        and ws["A1"].font.color.rgb.endswith("FFFFFF")
    )
    assert ws.freeze_panes == "A2" and ws.auto_filter.ref == "A1:I4"
    assert [ws.cell(row=r, column=1).value for r in range(2, 5)] == ["TC-01", "TC-02", "TC-18"]
    ss = wb["Summary"]
    col_a = [ss.cell(row=r, column=1).value for r in range(1, ss.max_row + 1)]
    for label in ref["summary_rows"]:
        assert label in col_a, label
    assert ss["A1"].value == "Test Case Matrix — X" and ss["A1"].font.size == 14
    assert ss["B3"].value == "TDD-ORD-002" and ss["B4"].value == "EPIC-002"
    totals = {
        ss.cell(row=r, column=1).value: ss.cell(row=r, column=2).value
        for r in range(1, ss.max_row + 1)
    }
    assert totals["Automated (Yes)"] == 1 and totals["Manual"] == 1 and totals["Planned"] == 1
    assert totals["Total Test Cases"] == 3
    assert (
        totals["Functional"] == 1
        and totals["Functional / Compensation"] == 1
        and totals["Reliability"] == 1
    )
    assert totals["Edge case"] == 0
    assert col_a[-2:] == ["n1", "n2"]
    # Reader round trip
    back = read_test_case_rows(data)
    assert [r.tc_id for r in back] == ["TC-01", "TC-02", "TC-18"] and back[
        1
    ].automated == "Manual (chaos test)"


def test_read_reference_matrix() -> None:
    rows = read_test_case_rows(XLSX_REF.read_bytes())
    assert [r.tc_id for r in rows][:3] == ["TC-01", "TC-02", "TC-03"] and len(rows) == 17
    assert rows[4].title.startswith("Provisioning fails") and rows[4].type == "Functional"


def test_preview_test_case_rows_merge_and_fallback() -> None:
    doc = parse_impact(_md("BCR-001-impact-analysis.md"))
    plain = preview_test_case_rows(doc)
    ids = [r.tc_id for r in plain]
    assert ids == sorted(set(ids), key=lambda s: (len(s), s)) and "TC-06" in ids and "TC-10" in ids
    assert all(r.preconditions == "" for r in plain) and all(r.notes for r in plain)
    existing = read_test_case_rows(XLSX_REF.read_bytes())
    merged = preview_test_case_rows(doc, existing)
    assert [r.tc_id for r in merged] == ids
    tc06 = next(r for r in merged if r.tc_id == "TC-06")
    assert (
        tc06.title.startswith("Dispatch fails after provisioning")
        and tc06.type == "Functional / Compensation"
    )
    assert tc06.preconditions and tc06.steps and tc06.expected and tc06.automated == "Yes"
    assert "modify" in tc06.notes
    data = export_test_case_preview(
        doc,
        existing=existing,
        label="Approved v6",
        linked_tdd="TDD-ORD-002",
        linked_epic="EPIC-002",
    )
    wb = load_workbook(io.BytesIO(data))
    ss = wb["Summary"]
    assert ss["A1"].value.endswith("(preview)") and ss["B3"].value == "TDD-ORD-002"
    assert wb["Test Cases"].max_row == len(ids) + 1


# --------------------------------------------------------------------------- #
# Change-request level: export_artifact, draft labels, the zip bundle
# --------------------------------------------------------------------------- #


def _cr(*, approved: bool = True) -> ChangeRequest:
    cr = ChangeRequest(
        kb_id="kb", title="Partial Shipment Support for Multi-Line Orders", document_text="x"
    )
    cr.bcr_meta.doc_id = "BCR-001"
    cr.ids.epic_id = "EPIC-002"
    cr.ids.tdd_id = "TDD-ORD-002"
    cr.ids.story_ids = [f"US-{n:03d}" for n in range(8, 16)]
    for kind, name in (
        (ArtifactKind.IMPACT, "BCR-001-impact-analysis.md"),
        (ArtifactKind.EPIC, "EPIC-002.md"),
        (ArtifactKind.STORIES, "US-008-015-stories.md"),
        (ArtifactKind.TDD, "TDD-ORD-002.md"),
    ):
        artifact: Artifact = cr.artifacts.get(kind)
        artifact.add_version(_md(name), VersionSource.LLM_DRAFT, "draft")
        if approved:
            artifact.status = ArtifactStatus.APPROVED
            artifact.approved_at = datetime(2026, 8, 18, tzinfo=UTC)
    return cr


def test_export_label_and_draft_suffix() -> None:
    cr = _cr(approved=False)
    assert export_label(cr.artifacts.impact) == "DRAFT v1 — not approved"
    out = export_artifact(cr, "impact", "docx")
    assert out.filename == "Impact-Analysis-BCR-001-DRAFT.docx"
    doc = _doc(out.data)
    assert doc.paragraphs[1].text.endswith("— DRAFT v1 — not approved")
    md = export_artifact(cr, "epic", "md")
    assert md.filename == "EPIC-002-DRAFT.md" and md.data.decode("utf-8").startswith("# EPIC-002")
    approved = _cr()
    assert export_label(approved.artifacts.impact) == "Approved v1 (2026-08-18)"
    assert export_artifact(approved, "impact", "docx").filename == "Impact-Analysis-BCR-001.docx"


def test_export_artifact_formats_and_errors() -> None:
    cr = _cr()
    stories = export_artifact(cr, ArtifactKind.STORIES, "docx")
    assert (
        stories.media_type == "application/zip" and stories.filename == "EPIC-002-user-stories.zip"
    )
    with zipfile.ZipFile(io.BytesIO(stories.data)) as zf:
        names = zf.namelist()
    assert len(names) == 8 and names[0].startswith("US-008-") and names[-1].startswith("US-015-")
    xlsx = export_artifact(cr, "impact", "xlsx")
    assert xlsx.filename == "TC-preview-BCR-001.xlsx"
    assert load_workbook(io.BytesIO(xlsx.data)).sheetnames == ["Test Cases", "Summary"]
    with pytest.raises(ValueError, match="Only the impact analysis"):
        export_artifact(cr, "epic", "xlsx")
    with pytest.raises(ValueError, match="Unsupported export format"):
        export_artifact(cr, "epic", "pdf")
    empty = ChangeRequest(kb_id="kb", title="t", document_text="x")
    with pytest.raises(ValueError, match="not been drafted"):
        export_artifact(empty, "impact", "docx")


def test_bundle_manifest() -> None:
    cr = _cr()
    data = export_change_request(cr)
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = zf.namelist()
        manifest = zf.read("MANIFEST.txt").decode("utf-8")
        assert zf.testzip() is None
    docx = [n for n in names if n.endswith(".docx")]
    assert "Impact-Analysis-BCR-001.docx" in docx
    assert any(n.startswith("EPIC-002-") for n in docx)
    assert sum(1 for n in docx if n.startswith("US-0")) == 8
    assert "TDD-ORD-002-orderworkflow-temporal-implementation.docx" in docx
    assert "TC-preview-BCR-001.xlsx" in names
    assert {
        "markdown/BCR-001-impact-analysis.md",
        "markdown/EPIC-002.md",
        "markdown/EPIC-002-user-stories.md",
        "markdown/TDD-ORD-002.md",
        "MANIFEST.txt",
    } <= set(names)
    assert "Change request: Partial Shipment Support" in manifest
    assert "Impact-Analysis-BCR-001.docx: impact — Approved v1" in manifest
    # Deterministic bytes.
    assert export_change_request(cr) == data
    # A partially drafted CR skips missing artifacts and says so.
    partial = ChangeRequest(kb_id="kb", title="t", document_text="x")
    partial.bcr_meta.doc_id = "BCR-009"
    partial.artifacts.impact.add_version(
        _md("BCR-001-impact-analysis.md"), VersionSource.LLM_DRAFT, "d"
    )
    with zipfile.ZipFile(io.BytesIO(export_change_request(partial))) as zf:
        assert "Impact-Analysis-BCR-001-DRAFT.docx" in zf.namelist()
        assert "tdd: not drafted — skipped" in zf.read("MANIFEST.txt").decode("utf-8")
