"""KgService: create → index → retrieve / impact / search / catalog / files, plus enrichment."""

from __future__ import annotations

import json
import shutil
from pathlib import Path

import pytest
from docx import Document

from workflow_compiler.exceptions import CompilationError, StateNotFoundError
from workflow_compiler.kg import (
    FileKnowledgeBaseStore,
    InMemoryKnowledgeBaseStore,
    KgService,
    KnowledgeBase,
    validate_kb_id,
)
from workflow_compiler.kg.ingest import zip_folder
from workflow_compiler.kg.llm_bridge import ProviderJsonClient
from workflow_compiler.llm.providers.mock import MockProvider

FIXTURE = Path(__file__).parent / "fixtures" / "kb_mini"


def build_corpus(dest: Path) -> Path:
    """Copy the fixture and add a docx (generated here so no binary lives in git)."""
    shutil.copytree(FIXTURE, dest)
    doc = Document()
    doc.add_heading("User Story US-005 — Complete order", level=1)
    doc.add_paragraph("As an operator I want the order marked complete after dispatch.")
    doc.add_paragraph("Acceptance: TC-03 verifies completion after dispatch_order succeeds.")
    doc.save(dest / "docs" / "US-005-complete-order.docx")
    return dest


@pytest.fixture
def corpus(tmp_path: Path) -> Path:
    return build_corpus(tmp_path / "kb_mini")


@pytest.fixture
def service(tmp_path: Path) -> KgService:
    return KgService(InMemoryKnowledgeBaseStore(tmp_path / "state"))


@pytest.fixture
async def ready_kb(service: KgService, corpus: Path) -> KnowledgeBase:
    kb = await service.create_from_zip("mini", zip_folder(corpus), owner_id="u1", filename="m.zip")
    return await service.index(kb.kb_id)


# ------------------------------------------------------------------ lifecycle


async def test_create_from_zip_extracts_and_records(service: KgService, corpus: Path) -> None:
    kb = await service.create_from_zip("mini", zip_folder(corpus), owner_id="u1", filename="m.zip")
    assert kb.status == "ingesting"
    assert kb.owner_id == "u1"
    assert kb.source.kind == "zip" and kb.source.filename == "m.zip"
    assert kb.stats.files == 8
    assert (service.corpus_dir(kb.kb_id) / "docs" / "BRD-order.md").is_file()
    assert any("kb_mini" in w for w in kb.warnings)  # top-level folder stripped
    assert Path(kb.root_dir) == service.store.kb_dir(kb.kb_id)


async def test_index_builds_graph_with_posix_ids_and_catalog(ready_kb: KnowledgeBase) -> None:
    kb = ready_kb
    assert kb.status == "ready"
    assert kb.indexed_at is not None
    assert kb.llm_enriched is False and kb.provider_used is None
    by_type = kb.stats.by_type
    assert by_type["Document"] >= 4  # 2 md + mmd + docx (README belongs to the repo node)
    assert by_type["Module"] == 3
    assert by_type["Function"] >= 4  # the four activities (methods hang off the class)
    assert by_type["Class"] == 1
    assert by_type["Chunk"] > 0
    assert kb.stats.edges > 0
    # ids are relative to corpus/ and use POSIX separators on every OS
    catalog = kb.catalog
    assert "US-001" in catalog.stories and "US-005" in catalog.stories
    assert "TC-01" in catalog.test_cases and "TC-03" in catalog.test_cases
    assert "BR-01" in catalog.requirements
    graph_json = json.loads(
        (Path(kb.root_dir) / ".contexthub" / "graph.json").read_text(encoding="utf-8")
    )
    ids = [n["id"] for n in graph_json["nodes"]]
    assert "mod:src/orders/workflow.py" in ids
    assert not any("\\" in i for i in ids)
    for node in graph_json["nodes"]:
        for key in ("path", "repo_path", "file", "extract_path"):
            value = (node.get("metadata") or {}).get(key)
            if isinstance(value, str) and value[1:3] != ":/":
                assert "\\" not in value, (node["id"], key, value)


async def test_index_failure_is_recorded(service: KgService, corpus: Path) -> None:
    kb = await service.create_from_zip("mini", zip_folder(corpus))
    shutil.rmtree(service.corpus_dir(kb.kb_id))
    with pytest.raises(CompilationError):
        await service.index(kb.kb_id)


async def test_list_delete_and_missing(service: KgService, corpus: Path) -> None:
    kb = await service.create_from_zip("mini", zip_folder(corpus))
    assert [k.kb_id for k in await service.list_all()] == [kb.kb_id]
    await service.delete(kb.kb_id)
    assert await service.list_all() == []
    assert not service.store.kb_dir(kb.kb_id).exists()
    with pytest.raises(StateNotFoundError):
        await service.get(kb.kb_id)
    with pytest.raises(StateNotFoundError):
        await service.retrieve("../etc", "x")


async def test_retrieve_before_index_is_a_clean_error(service: KgService, corpus: Path) -> None:
    kb = await service.create_from_zip("mini", zip_folder(corpus))
    with pytest.raises(CompilationError, match="no graph yet"):
        await service.retrieve(kb.kb_id, "anything")


# -------------------------------------------------------------------- queries


async def test_retrieve_returns_grounded_sections(
    service: KgService, ready_kb: KnowledgeBase
) -> None:
    packet = await service.retrieve(
        ready_kb.kb_id, "how does dispatch compensate provisioning", budget=1500
    )
    assert packet.rendered
    assert packet.sections
    assert packet.total_tokens > 0
    paths = {f.path for f in packet.files}
    assert any(p.endswith(("workflow.py", "TDD-order.md", "BRD-order.md")) for p in paths)
    # at least one section dereferences a real line span
    assert any(s.start_line is not None and s.end_line is not None for s in packet.sections)
    assert 0.0 <= packet.coverage <= 1.0


async def test_search_ranks_the_workflow_module(
    service: KgService, ready_kb: KnowledgeBase
) -> None:
    hits = await service.search(ready_kb.kb_id, "release_provisioning compensation", k=5)
    assert hits
    assert all(h.score > 0 for h in hits)
    assert any("activities.py" in (h.path or "") or "workflow.py" in (h.path or "") for h in hits)


async def test_impact_bfs_is_deterministic(service: KgService, ready_kb: KnowledgeBase) -> None:
    rows = await service.impact(ready_kb.kb_id, ["mod:src/orders/activities.py"], max_hops=2)
    assert rows[0].node_id == "mod:src/orders/activities.py" and rows[0].hops == 0
    ids = [r.node_id for r in rows]
    assert "mod:src/orders/workflow.py" in ids  # IMPORTS edge
    assert any(i.startswith("fn:src/orders/activities.py:") for i in ids)  # CONTAINS downwards
    assert "repo:kb_mini" not in ids and not any(r.type == "Repository" for r in rows)
    again = await service.impact(ready_kb.kb_id, ["mod:src/orders/activities.py"], max_hops=2)
    assert [(r.node_id, r.hops, r.via) for r in again] == [(r.node_id, r.hops, r.via) for r in rows]
    # hops never decrease along the table
    assert [r.hops for r in rows] == sorted(r.hops for r in rows)


async def test_impact_resolves_terms_to_anchors(
    service: KgService, ready_kb: KnowledgeBase
) -> None:
    rows = await service.impact(ready_kb.kb_id, ["dispatch_order"], max_hops=1)
    assert rows and any(r.hops == 0 for r in rows)
    assert await service.impact(ready_kb.kb_id, ["   "], max_hops=1) == []


async def test_catalog_and_graph_summary(service: KgService, ready_kb: KnowledgeBase) -> None:
    catalog = await service.catalog(ready_kb.kb_id)
    assert catalog == ready_kb.catalog
    summary = await service.graph_summary(ready_kb.kb_id, top=5)
    assert summary.nodes == ready_kb.stats.nodes
    assert summary.by_type == ready_kb.stats.by_type
    assert len(summary.top_nodes) == 5
    assert all(n.type not in ("Repository", "Service") for n in summary.top_nodes)


async def test_files_listing_and_reading(service: KgService, ready_kb: KnowledgeBase) -> None:
    files = await service.list_files(ready_kb.kb_id)
    assert "src/orders/workflow.py" in files and "docs/US-005-complete-order.docx" in files
    py = await service.read_file(ready_kb.kb_id, "src/orders/workflow.py")
    assert "class OrderWorkflow" in py.text and py.extracted is False
    docx = await service.read_file(ready_kb.kb_id, "docs/US-005-complete-order.docx")
    assert docx.extracted is True and "TC-03" in docx.text
    with pytest.raises(StateNotFoundError):
        await service.read_file(ready_kb.kb_id, "../../etc/passwd")
    with pytest.raises(StateNotFoundError):
        await service.read_file(ready_kb.kb_id, "docs/nope.md")


# ---------------------------------------------------------------- enrichment


_ENRICH = {
    "summary": "Explains the order lifecycle and its saga compensation.",
    "topics": ["order lifecycle", "saga compensation"],
    "entities": ["OrderWorkflow", "dispatch_order"],
    "doc_type": "tdd",
}
_CLUSTER = {
    "processes": [
        {"id": "order-lifecycle", "name": "Order lifecycle", "summary": "s", "file_ids": []}
    ],
    "unassigned": [],
}


class EnrichmentMock(MockProvider):
    """Answers the per-file enrichment prompt or the clustering prompt by content."""

    async def complete(self, prompt: str, **kwargs: object) -> str:  # type: ignore[override]
        self.calls.append(("complete", prompt))
        payload = _CLUSTER if "cluster these" in prompt.lower() else _ENRICH
        return "```json\n" + json.dumps(payload) + "\n```"


async def test_index_with_enrichment_uses_the_provider_factory(
    tmp_path: Path, corpus: Path
) -> None:
    provider = EnrichmentMock()
    seen: list[tuple[str | None, str | None]] = []

    def factory(name: str | None, model: str | None) -> MockProvider:
        seen.append((name, model))
        return provider

    service = KgService(InMemoryKnowledgeBaseStore(tmp_path / "state"), factory)
    kb = await service.create_from_zip("mini", zip_folder(corpus))
    progress: list[tuple[str, int, int]] = []
    kb = await service.index(
        kb.kb_id,
        enrich=True,
        provider="nemotron",
        model="m",
        progress=lambda *a: progress.append(a),
    )
    assert seen == [("nemotron", "m")]
    assert kb.status == "ready" and kb.llm_enriched is True
    assert kb.provider_used == "mock" and kb.model_used == "m"
    assert provider.calls, "the provider was asked"
    assert kb.stats.by_type.get("DataArtifact", 0) >= 2  # topic + entity nodes were minted
    assert progress and progress[0][0] == "static ingest"
    assert any(p[0] == "process clustering" for p in progress)
    # the cache makes a re-index free of LLM calls
    calls_before = len(provider.calls)
    await service.reindex(kb.kb_id, enrich=True, provider="nemotron", model="m")
    assert len(provider.calls) == calls_before


async def test_enrichment_without_factory_is_an_error(service: KgService, corpus: Path) -> None:
    kb = await service.create_from_zip("mini", zip_folder(corpus))
    with pytest.raises(CompilationError, match="provider factory"):
        await service.index(kb.kb_id, enrich=True)


def test_provider_json_client_parses_fences_and_retries() -> None:
    provider = MockProvider(completions=["[1, 2]", "prose ```json {\"a\": 1} ```"])
    client = ProviderJsonClient(provider)  # no running loop → asyncio.run path
    out = client.chat_json([{"role": "system", "content": "s"}, {"role": "user", "content": "u"}])
    assert out == {"a": 1}
    assert client.calls == 2 and client.failures == 0
    assert provider.calls[0][1] == "u"


def test_provider_json_client_gives_up_after_retries() -> None:
    from workflow_compiler.kg.contexthub.bootstrap.llm import LlmError

    provider = MockProvider(default_completion="not json at all")
    client = ProviderJsonClient(provider)
    with pytest.raises(LlmError):
        client.chat_json([{"role": "user", "content": "u"}], label="x", retries=2)
    assert client.calls == 2 and client.failures == 1


# ------------------------------------------------------------------ the store


async def test_file_store_round_trip_and_id_validation(tmp_path: Path) -> None:
    store = FileKnowledgeBaseStore(tmp_path)
    kb = KnowledgeBase(name="n", owner_id="o")
    await store.save(kb)
    loaded = await store.load(kb.kb_id)
    assert loaded == kb
    assert await store.list_ids() == [kb.kb_id]
    assert await store.exists(kb.kb_id)
    assert store.kb_dir(kb.kb_id) == tmp_path.resolve() / "knowledge_bases" / kb.kb_id
    await store.delete(kb.kb_id)
    assert await store.list_ids() == []
    for bad in ("../x", "a/b", "a\\b", "", "C:x", "a b"):
        with pytest.raises(StateNotFoundError):
            validate_kb_id(bad)
        with pytest.raises(StateNotFoundError):
            await store.load(bad)
    assert validate_kb_id("kb_1-A") == "kb_1-A"


async def test_service_on_file_store(tmp_path: Path, corpus: Path) -> None:
    service = KgService(FileKnowledgeBaseStore(tmp_path / "state"))
    kb = await service.create_from_zip("mini", zip_folder(corpus))
    kb = await service.index(kb.kb_id)
    assert (tmp_path / "state" / "knowledge_bases" / f"{kb.kb_id}.json").is_file()
    kb_dir = tmp_path / "state" / "knowledge_bases" / kb.kb_id
    assert (kb_dir / ".contexthub" / "graph.json").is_file()
    packet = await service.retrieve(kb.kb_id, "validate order")
    assert packet.sections
